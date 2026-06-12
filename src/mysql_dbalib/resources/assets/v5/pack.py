#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Самораспаковывающийся исходник проекта «МебельОрг»."""

from __future__ import annotations

import argparse
import py_compile
import subprocess
import sys
from pathlib import Path


EMPTY_DIRS = [
    'app/assets',
    'app/assets/products',
    'database',
    'docs',
    'docs/screenshots',
    'tools',
]


FILES = {
    # ============================================================
    # BEGIN FILE: README.md
    # ============================================================
    'README.md': r'''# МебельОрг: демонстрационный экзамен, вариант 5

## Запуск
1. Откройте папку проекта.
2. Запустите `run.bat`.
3. При первом запуске приложение создаст базу и заполнит ее начальными данными.

По умолчанию используется SQLite, чтобы проект можно было сразу проверить без сервера.

## Проверка без GUI
```powershell
python tools\smoke_test.py
```

## Запуск с MySQL
Установите зависимости:

```powershell
pip install -r requirements.txt
```

Задайте переменные окружения:

```powershell
$env:DB_ENGINE="mysql"
$env:DB_HOST="localhost"
$env:DB_PORT="3306"
$env:DB_NAME="mebelorg"
$env:DB_USER="root"
$env:DB_PASSWORD="пароль"
python app\main.py
```

Для MySQL нужен пакет `mysql-connector-python`, в коде он импортируется как `mysql.connector`.

## Изображения
В переносимый вариант изображения не вшиты. Если на экзамене выдали ресурсы, скопируйте их так:

- `icon.png` и `icon.ico` в `app\assets`
- `picture.png` в `app\assets`
- фото товаров `1.jpg`, `2.jpg` и далее в `app\assets\products`

Если изображений нет, приложение создаст или покажет заглушку и продолжит работать.

## Тестовые пользователи
- Администратор: `94d5ous@gmail.com` / `uzWC67`
- Менеджер: `ptec8ym@yahoo.com` / `LdNyos`
- Авторизированный клиент: `yzls62@outlook.com` / `JlFRCZ`

## Основные файлы
- `app/main.py` - точка входа приложения.
- `database/schema_sqlite.sql` - SQL-скрипт создания SQLite-базы.
- `database/schema_mysql.sql` - SQL-скрипт создания MySQL-базы.
- `database/mebelorg.db` - создается автоматически при запуске в SQLite-режиме.
- `database/seed_sqlite.sql` и `database/seed_mysql.sql` - начальные данные.
- `database/import_report.txt` - отчет о мягкой очистке import-данных.
- `docs/ER-diagram.pdf` - ER-диаграмма.
- `docs/flowchart.pdf` - блок-схема алгоритма.
- `docs/screenshots.docx` - документ со скриншотами работы системы.
''',
    # ============================================================
    # END FILE: README.md
    # ============================================================

    # ============================================================
    # BEGIN FILE: requirements.txt
    # ============================================================
    'requirements.txt': r'''openpyxl
pillow
python-docx
reportlab
mysql-connector-python
''',
    # ============================================================
    # END FILE: requirements.txt
    # ============================================================

    # ============================================================
    # BEGIN FILE: run.bat
    # ============================================================
    'run.bat': r'''@echo off
chcp 65001 >nul
cd /d "%~dp0"
where python >nul 2>nul
if %errorlevel% equ 0 (
    python app\main.py
) else (
    py -3 app\main.py
)
if errorlevel 1 pause
''',
    # ============================================================
    # END FILE: run.bat
    # ============================================================

    # ============================================================
    # BEGIN FILE: ЭТАПЫ.md
    # ============================================================
    'ЭТАПЫ.md': r'''# Этапы подготовки решения

- [x] Прочитать текст задания и приложения варианта 5.
- [x] Подготовить структуру проекта: `app`, `database`, `docs`, `tools`.
- [x] Создать SQLite-схему в 3НФ с первичными и внешними ключами.
- [x] Импортировать товары, пользователей, заказы и пункты выдачи.
- [x] Реализовать мягкую очистку import-данных и отчет о ней.
- [x] Реализовать авторизацию, роли и выход на экран входа.
- [x] Реализовать список товаров для гостя, клиента, менеджера и администратора.
- [x] Добавить поиск, сортировку и фильтр по скидке для менеджера и администратора.
- [x] Реализовать добавление, редактирование и удаление товаров для администратора.
- [x] Реализовать список заказов для менеджера и администратора.
- [x] Реализовать добавление, редактирование и удаление заказов для администратора.
- [x] Сгенерировать ER-диаграмму, блок-схему и документ со скриншотами.
- [x] Выполнить smoke-тест и финальную проверку запуска.
''',
    # ============================================================
    # END FILE: ЭТАПЫ.md
    # ============================================================

    # ============================================================
    # BEGIN FILE: app/__init__.py
    # ============================================================
    'app/__init__.py': r'''"""Приложение ООО «МебельОрг» для демонстрационного экзамена."""
''',
    # ============================================================
    # END FILE: app/__init__.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: app/config.py
    # ============================================================
    'app/config.py': r'''import os
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[1]
APP_DIR = PROJECT_ROOT / "app"
DATABASE_DIR = PROJECT_ROOT / "database"
DOCS_DIR = PROJECT_ROOT / "docs"
ASSETS_DIR = APP_DIR / "assets"
PRODUCT_IMAGES_DIR = ASSETS_DIR / "products"

DATABASE_PATH = DATABASE_DIR / "mebelorg.db"
SCHEMA_PATH = DATABASE_DIR / "schema.sql"
SCHEMA_SQLITE_PATH = DATABASE_DIR / "schema_sqlite.sql"
SEED_SQLITE_PATH = DATABASE_DIR / "seed_sqlite.sql"
SCHEMA_MYSQL_PATH = DATABASE_DIR / "schema_mysql.sql"
SEED_MYSQL_PATH = DATABASE_DIR / "seed_mysql.sql"
IMPORT_RAR_PATH = PROJECT_ROOT / "Модуль 1" / "Прил_2_ОЗ_КОД 09.02.07-2-2026-М1.rar"
IMPORT_SOURCE_DIR = ASSETS_DIR / "import"

DB_ENGINE = os.environ.get("DB_ENGINE", "sqlite").strip().lower()
DB_HOST = os.environ.get("DB_HOST", "localhost")
DB_PORT = int(os.environ.get("DB_PORT", "3306"))
DB_NAME = os.environ.get("DB_NAME", "mebelorg")
DB_USER = os.environ.get("DB_USER", "root")
DB_PASSWORD = os.environ.get("DB_PASSWORD", "")

APP_TITLE = "МебельОрг"
FONT_FAMILY = "Calibri"

COLOR_WHITE = "#FFFFFF"
COLOR_SECONDARY = "#00FFFF"
COLOR_ACCENT = "#0000FF"
COLOR_DISCOUNT = "#008080"
COLOR_OUT_OF_STOCK = "#D9D9D9"
COLOR_PANEL = "#EFFFFF"

LOGO_PATH = ASSETS_DIR / "icon.png"
ICON_PATH = ASSETS_DIR / "icon.ico"
PLACEHOLDER_PATH = ASSETS_DIR / "picture.png"

ROLE_GUEST = "Гость"
ROLE_CLIENT = "Авторизированный клиент"
ROLE_MANAGER = "Менеджер"
ROLE_ADMIN = "Администратор"

DISCOUNT_RANGES = {
    "Все диапазоны": None,
    "0-10,99%": (0, 10.99),
    "11-14,99%": (11, 14.99),
    "15% и более": (15, 100),
}
''',
    # ============================================================
    # END FILE: app/config.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: app/database.py
    # ============================================================
    'app/database.py': r'''from __future__ import annotations

import sqlite3
from pathlib import Path

from config import (
    DATABASE_DIR,
    DATABASE_PATH,
    DB_ENGINE,
    DB_HOST,
    DB_NAME,
    DB_PASSWORD,
    DB_PORT,
    DB_USER,
    IMPORT_RAR_PATH,
    IMPORT_SOURCE_DIR,
    SCHEMA_MYSQL_PATH,
    SCHEMA_PATH,
    SCHEMA_SQLITE_PATH,
    SEED_MYSQL_PATH,
    SEED_SQLITE_PATH,
)


class ResultRow(dict):
    def __getitem__(self, key):
        if isinstance(key, int):
            return list(self.values())[key]

        return super().__getitem__(key)


class DatabaseCursor:
    def __init__(self, cursor) -> None:
        self.cursor = cursor
        self.lastrowid = getattr(cursor, "lastrowid", None)

    def fetchone(self):
        row = self.cursor.fetchone()
        return self._wrap_row(row)

    def fetchall(self):
        return [self._wrap_row(row) for row in self.cursor.fetchall()]

    def _wrap_row(self, row):
        if row is None:
            return None

        if isinstance(row, dict):
            return ResultRow(row)

        return ResultRow(dict(row))


class DatabaseConnection:
    def __init__(self) -> None:
        self.engine = DB_ENGINE
        self.connection = self._connect()

    def __enter__(self):
        return self

    def __exit__(self, exc_type, _exc_value, _traceback) -> None:
        if exc_type:
            self.connection.rollback()
        else:
            self.connection.commit()

        self.connection.close()

    def execute(self, sql: str, params: tuple | list | None = None) -> DatabaseCursor:
        cursor = self._cursor()
        cursor.execute(self._prepare_sql(sql), tuple(params or ()))
        return DatabaseCursor(cursor)

    def executemany(self, sql: str, params: list[tuple]) -> None:
        cursor = self._cursor()
        cursor.executemany(self._prepare_sql(sql), params)

    def executescript(self, script: str) -> None:
        if self.engine == "sqlite":
            self.connection.executescript(script)
            return

        cursor = self._cursor()
        for statement in split_sql_script(script):
            cursor.execute(statement)

    def iterdump(self):
        if self.engine != "sqlite":
            raise RuntimeError("SQL-дамп через iterdump доступен только для SQLite.")

        return self.connection.iterdump()

    def commit(self) -> None:
        self.connection.commit()

    def _connect(self):
        if self.engine == "sqlite":
            DATABASE_DIR.mkdir(parents=True, exist_ok=True)
            connection = sqlite3.connect(DATABASE_PATH)
            connection.row_factory = sqlite3.Row
            connection.execute("PRAGMA foreign_keys = ON")
            return connection

        if self.engine == "mysql":
            return connect_mysql()

        raise RuntimeError("DB_ENGINE должен быть sqlite или mysql.")

    def _cursor(self):
        if self.engine == "mysql":
            return self.connection.cursor(dictionary=True)

        return self.connection.cursor()

    def _prepare_sql(self, sql: str) -> str:
        if self.engine == "mysql":
            return sql.replace("?", "%s")

        return sql


def connect_mysql():
    try:
        import mysql.connector
    except ModuleNotFoundError as error:
        raise RuntimeError(
            "Для режима MySQL установите пакет: pip install mysql-connector-python"
        ) from error

    return mysql.connector.connect(
        host=DB_HOST,
        port=DB_PORT,
        user=DB_USER,
        password=DB_PASSWORD,
        database=DB_NAME,
        charset="utf8mb4",
        use_unicode=True,
    )


def create_mysql_database_if_needed() -> None:
    try:
        import mysql.connector
    except ModuleNotFoundError as error:
        raise RuntimeError(
            "Для режима MySQL установите пакет: pip install mysql-connector-python"
        ) from error

    connection = mysql.connector.connect(
        host=DB_HOST,
        port=DB_PORT,
        user=DB_USER,
        password=DB_PASSWORD,
        charset="utf8mb4",
        use_unicode=True,
    )
    try:
        cursor = connection.cursor()
        cursor.execute(
            f"CREATE DATABASE IF NOT EXISTS `{DB_NAME}` CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci"
        )
        connection.commit()
    finally:
        connection.close()


def get_connection() -> DatabaseConnection:
    return DatabaseConnection()


def initialize_database(reset: bool = False) -> None:
    DATABASE_DIR.mkdir(parents=True, exist_ok=True)

    if reset:
        reset_database_storage()

    if DB_ENGINE == "mysql":
        create_mysql_database_if_needed()

    schema_path = get_schema_path()
    seed_path = get_seed_path()

    with get_connection() as connection:
        schema = schema_path.read_text(encoding="utf-8")
        connection.executescript(schema)

        product_count = connection.execute("SELECT COUNT(*) FROM products").fetchone()[0]
        if product_count == 0:
            if seed_path.exists():
                seed = seed_path.read_text(encoding="utf-8")
                connection.executescript(seed)
            elif DB_ENGINE == "sqlite" and (IMPORT_SOURCE_DIR.exists() or IMPORT_RAR_PATH.exists()):
                from import_data import import_all_data

                import_all_data(connection)
            else:
                raise RuntimeError("Не найден файл с начальными данными для базы.")


def rebuild_database() -> None:
    initialize_database(reset=True)


def reset_database_storage() -> None:
    if DB_ENGINE == "sqlite":
        if DATABASE_PATH.exists():
            DATABASE_PATH.unlink()
        return

    if DB_ENGINE == "mysql":
        create_mysql_database_if_needed()
        with get_connection() as connection:
            connection.executescript(get_drop_mysql_script())
        return

    raise RuntimeError("DB_ENGINE должен быть sqlite или mysql.")


def create_sql_dump(output_path: Path) -> None:
    with get_connection() as connection:
        dump_text = "\n".join(connection.iterdump())

    output_path.write_text(dump_text, encoding="utf-8")


def get_schema_path() -> Path:
    if DB_ENGINE == "mysql":
        return SCHEMA_MYSQL_PATH

    if SCHEMA_SQLITE_PATH.exists():
        return SCHEMA_SQLITE_PATH

    return SCHEMA_PATH


def get_seed_path() -> Path:
    if DB_ENGINE == "mysql":
        return SEED_MYSQL_PATH

    return SEED_SQLITE_PATH


def get_drop_mysql_script() -> str:
    return """
    SET FOREIGN_KEY_CHECKS = 0;
    DROP TABLE IF EXISTS order_items;
    DROP TABLE IF EXISTS orders;
    DROP TABLE IF EXISTS products;
    DROP TABLE IF EXISTS users;
    DROP TABLE IF EXISTS roles;
    DROP TABLE IF EXISTS categories;
    DROP TABLE IF EXISTS suppliers;
    DROP TABLE IF EXISTS manufacturers;
    DROP TABLE IF EXISTS units;
    DROP TABLE IF EXISTS pickup_points;
    DROP TABLE IF EXISTS order_statuses;
    SET FOREIGN_KEY_CHECKS = 1;
    """


def split_sql_script(script: str) -> list[str]:
    cleaned_lines = []
    for line in script.splitlines():
        stripped = line.strip()
        if stripped.startswith("--") or stripped.startswith("PRAGMA"):
            continue
        cleaned_lines.append(line)

    cleaned_script = "\n".join(cleaned_lines)
    statements = [statement.strip() for statement in cleaned_script.split(";")]
    return [statement for statement in statements if statement]
''',
    # ============================================================
    # END FILE: app/database.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: app/import_data.py
    # ============================================================
    'app/import_data.py': r'''from __future__ import annotations

import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Iterable

from openpyxl import load_workbook
from config import (
    ASSETS_DIR,
    DATABASE_DIR,
    IMPORT_RAR_PATH,
    IMPORT_SOURCE_DIR,
    PRODUCT_IMAGES_DIR,
    ROLE_ADMIN,
    ROLE_CLIENT,
    ROLE_MANAGER,
)


ROLE_NAMES = [ROLE_ADMIN, ROLE_MANAGER, ROLE_CLIENT]
IMPORT_REPORT_PATH = DATABASE_DIR / "import_report.txt"
MAX_PRODUCT_IMAGE_SIZE = (300, 200)


def import_all_data(connection) -> None:
    source_dir = _ensure_import_source()
    _copy_common_assets(source_dir)
    PRODUCT_IMAGES_DIR.mkdir(parents=True, exist_ok=True)

    report: list[str] = []
    workbooks = _load_workbooks(source_dir)

    _insert_roles(connection)
    _import_users(connection, workbooks["users"])
    _import_pickup_points(connection, workbooks["pickup_points"])
    _import_products(connection, workbooks["products"], source_dir, report)
    _import_orders(connection, workbooks["orders"], report)
    connection.commit()
    _write_import_report(report)


def _ensure_import_source() -> Path:
    if IMPORT_SOURCE_DIR.exists():
        return IMPORT_SOURCE_DIR

    if not IMPORT_RAR_PATH.exists():
        raise RuntimeError("Не найден архив с исходными import-данными.")

    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        ["tar", "-xf", str(IMPORT_RAR_PATH), "-C", str(ASSETS_DIR)],
        check=True,
        capture_output=True,
        text=True,
    )
    return IMPORT_SOURCE_DIR


def _copy_common_assets(source_dir: Path) -> None:
    for name in ("picture.png", "icon.ico", "icon.jpg", "icon.png"):
        source = source_dir / name
        if source.exists():
            shutil.copy2(source, ASSETS_DIR / name)


def _load_workbooks(source_dir: Path) -> dict[str, list[tuple]]:
    mapping = {
        "products": source_dir / "Tovar.xlsx",
        "users": source_dir / "user_import.xlsx",
        "orders": source_dir / "Заказ_import.xlsx",
        "pickup_points": source_dir / "Пункты выдачи_import.xlsx",
    }
    result = {}

    for key, path in mapping.items():
        workbook = load_workbook(path, data_only=True, read_only=True)
        sheet = workbook.active
        result[key] = list(sheet.iter_rows(values_only=True))
        workbook.close()

    return result


def _insert_roles(connection) -> None:
    for role_name in ROLE_NAMES:
        connection.execute(
            "INSERT OR IGNORE INTO roles(name) VALUES (?)",
            (role_name,),
        )


def _import_users(connection, rows: list[tuple]) -> None:
    for role_name, full_name, login, password, *_ in rows[1:]:
        if not role_name or not full_name or not login or not password:
            continue

        role_id = _get_or_create(connection, "roles", role_name)
        connection.execute(
            """
            INSERT OR IGNORE INTO users(role_id, full_name, login, password)
            VALUES (?, ?, ?, ?)
            """,
            (role_id, _clean_text(full_name), _clean_text(login), _clean_text(password)),
        )


def _import_pickup_points(connection, rows: list[tuple]) -> None:
    for row in rows:
        address = _clean_text(row[0] if row else None)
        if not address:
            continue

        connection.execute(
            "INSERT OR IGNORE INTO pickup_points(address) VALUES (?)",
            (address,),
        )


def _import_products(connection, rows: list[tuple], source_dir: Path, report: list[str]) -> None:
    headers = list(rows[0])
    indexes = {name: index for index, name in enumerate(headers)}

    for row in rows[1:]:
        article = _clean_text(_cell(row, indexes, "Артикул"))
        if not article:
            continue

        unit_id = _get_or_create(connection, "units", _cell(row, indexes, "Единица измерения"))
        supplier_id = _get_or_create(connection, "suppliers", _cell(row, indexes, "Поставщик"))
        manufacturer_id = _get_or_create(connection, "manufacturers", _cell(row, indexes, "Производитель"))
        category_id = _get_or_create(connection, "categories", _cell(row, indexes, "Категория товара"))
        photo_path = _prepare_product_image(
            _cell(row, indexes, "Фото"),
            article,
            source_dir,
            report,
        )

        connection.execute(
            """
            INSERT OR IGNORE INTO products(
                article, name, unit_id, price, supplier_id, manufacturer_id,
                category_id, discount, stock_quantity, description, photo_path
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                article,
                _clean_text(_cell(row, indexes, "Наименование товара")),
                unit_id,
                _to_float(_cell(row, indexes, "Цена")),
                supplier_id,
                manufacturer_id,
                category_id,
                _to_float(_cell(row, indexes, "Действующая скидка")),
                _to_int(_cell(row, indexes, "Кол-во на складе")),
                _clean_text(_cell(row, indexes, "Описание товара")),
                photo_path,
            ),
        )


def _import_orders(connection, rows: list[tuple], report: list[str]) -> None:
    headers = list(rows[0])
    indexes = {name: index for index, name in enumerate(headers) if name}

    for row in rows[1:]:
        order_id = _cell(row, indexes, "Номер заказа")
        if order_id is None:
            continue

        status_id = _get_or_create(connection, "order_statuses", _cell(row, indexes, "Статус заказа"))
        pickup_point_id = _resolve_pickup_point(connection, _cell(row, indexes, "Адрес пункта выдачи"))
        client_name = _clean_text(_cell(row, indexes, "ФИО авторизированного клиента"))
        client_user_id = _find_user_by_name(connection, client_name)
        order_date = _parse_excel_date(_cell(row, indexes, "Дата заказа"), int(order_id), report)
        delivery_date = _parse_excel_date(_cell(row, indexes, "Дата доставки"), int(order_id), report)
        receive_code = _clean_text(_cell(row, indexes, "Код для получения"))

        connection.execute(
            """
            INSERT OR IGNORE INTO orders(
                id, pickup_point_id, client_user_id, client_name, receive_code,
                status_id, order_date, delivery_date
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                int(order_id),
                pickup_point_id,
                client_user_id,
                client_name,
                receive_code,
                status_id,
                order_date,
                delivery_date,
            ),
        )

        for article, quantity in _parse_order_items(_cell(row, indexes, "Артикул заказа")):
            product_id = _find_product_by_article(connection, article)
            if product_id is None:
                report.append(
                    f"Заказ {int(order_id)}: позиция {article} x {quantity} пропущена, "
                    "так как такого артикула нет в Tovar.xlsx."
                )
                continue

            connection.execute(
                """
                INSERT OR IGNORE INTO order_items(order_id, product_id, quantity)
                VALUES (?, ?, ?)
                """,
                (int(order_id), product_id, quantity),
            )


def _prepare_product_image(value, article: str, source_dir: Path, report: list[str]) -> str | None:
    try:
        from PIL import Image
    except ModuleNotFoundError as error:
        raise RuntimeError("Для импорта и сжатия изображений установите пакет Pillow.") from error

    photo_name = _clean_text(value)
    if not photo_name:
        return None

    source = source_dir / photo_name
    if not source.exists():
        # Если расширение в Excel отличается от реального файла, ищем фото по имени.
        candidates = list(source_dir.glob(f"{Path(photo_name).stem}.*"))
        image_candidates = [path for path in candidates if path.suffix.lower() in {".jpg", ".jpeg", ".png"}]
        source = image_candidates[0] if image_candidates else source

    if not source.exists():
        report.append(f"Товар {article}: изображение {photo_name} не найдено, используется заглушка.")
        return None

    destination = PRODUCT_IMAGES_DIR / f"{source.stem}{source.suffix.lower()}"
    with Image.open(source) as image:
        image.thumbnail(MAX_PRODUCT_IMAGE_SIZE)
        save_image = image.convert("RGB") if destination.suffix.lower() in {".jpg", ".jpeg"} else image
        save_image.save(destination)

    if source.name != photo_name:
        report.append(f"Товар {article}: фото {photo_name} заменено найденным файлом {source.name}.")

    return f"assets/products/{destination.name}"


def _cell(row: tuple, indexes: dict[str, int], column_name: str):
    index = indexes[column_name]
    return row[index] if index < len(row) else None


def _clean_text(value) -> str:
    if value is None:
        return ""

    return str(value).strip()


def _to_float(value) -> float:
    if value is None or value == "":
        return 0.0

    return float(str(value).replace(",", "."))


def _to_int(value) -> int:
    if value is None or value == "":
        return 0

    return int(float(str(value).replace(",", ".")))


def _get_or_create(connection, table_name: str, name_value) -> int:
    name = _clean_text(name_value)
    cursor = connection.execute(f"SELECT id FROM {table_name} WHERE name = ?", (name,))
    row = cursor.fetchone()

    if row:
        return row["id"]

    cursor = connection.execute(f"INSERT INTO {table_name}(name) VALUES (?)", (name,))
    return cursor.lastrowid


def _parse_order_items(value) -> Iterable[tuple[str, int]]:
    text = _clean_text(value)
    if not text:
        return []

    parts = [part.strip() for part in re.split(r"[,;\n]+", text) if part.strip()]
    items = []

    for index in range(0, len(parts), 2):
        if index + 1 >= len(parts):
            continue

        article = parts[index]
        quantity = _to_int(parts[index + 1])
        if article and quantity > 0:
            items.append((article, quantity))

    return items


def _parse_excel_date(value, order_id: int, report: list[str]) -> str | None:
    if value is None or value == "":
        return None

    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")

    text = str(value).strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(text, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue

    if re.match(r"^\d{2}\.\d{2}\.\d{4}$", text):
        report.append(f"Заказ {order_id}: некорректная дата {text} заменена пустым значением.")
        return None

    return text


def _resolve_pickup_point(connection, value) -> int:
    if isinstance(value, int):
        cursor = connection.execute("SELECT id FROM pickup_points WHERE id = ?", (value,))
        row = cursor.fetchone()
        if row:
            return row["id"]

    address = _clean_text(value)
    cursor = connection.execute("SELECT id FROM pickup_points WHERE address = ?", (address,))
    row = cursor.fetchone()
    if row:
        return row["id"]

    cursor = connection.execute("INSERT INTO pickup_points(address) VALUES (?)", (address or "Не указан"))
    return cursor.lastrowid


def _find_user_by_name(connection, full_name: str) -> int | None:
    cursor = connection.execute(
        """
        SELECT users.id
        FROM users
        JOIN roles ON roles.id = users.role_id
        WHERE users.full_name = ? AND roles.name = ?
        """,
        (full_name, ROLE_CLIENT),
    )
    row = cursor.fetchone()
    return row["id"] if row else None


def _find_product_by_article(connection, article: str) -> int | None:
    cursor = connection.execute("SELECT id FROM products WHERE article = ?", (article,))
    row = cursor.fetchone()
    return row["id"] if row else None


def _write_import_report(report: list[str]) -> None:
    DATABASE_DIR.mkdir(parents=True, exist_ok=True)
    lines = ["Отчет о мягкой очистке import-данных", ""]
    if report:
        lines.extend(f"- {message}" for message in report)
    else:
        lines.append("- Замечаний при импорте не найдено.")

    IMPORT_REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")
''',
    # ============================================================
    # END FILE: app/import_data.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: app/main.py
    # ============================================================
    'app/main.py': r'''import tkinter as tk
from tkinter import messagebox

from config import APP_TITLE
from database import initialize_database
from ui import Application


def main() -> None:
    try:
        initialize_database()
    except Exception as error:
        messagebox.showerror(
            "Ошибка запуска",
            f"Не удалось подготовить базу данных.\n\n{error}",
        )
        return

    root = tk.Tk()
    root.title(APP_TITLE)
    app = Application(root)
    app.show_login()
    root.mainloop()


if __name__ == "__main__":
    main()
''',
    # ============================================================
    # END FILE: app/main.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: app/repositories.py
    # ============================================================
    'app/repositories.py': r'''from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from config import ASSETS_DIR, DB_ENGINE, DISCOUNT_RANGES
from database import get_connection


@dataclass
class CurrentUser:
    id: int | None
    full_name: str
    role_name: str


class AuthRepository:
    def authenticate(self, login: str, password: str) -> CurrentUser | None:
        with get_connection() as connection:
            row = connection.execute(
                """
                SELECT users.id, users.full_name, roles.name AS role_name
                FROM users
                JOIN roles ON roles.id = users.role_id
                WHERE users.login = ? AND users.password = ?
                """,
                (login.strip(), password.strip()),
            ).fetchone()

        if not row:
            return None

        return CurrentUser(row["id"], row["full_name"], row["role_name"])


class LookupRepository:
    def get_values(self, table_name: str) -> list[dict[str, Any]]:
        with get_connection() as connection:
            rows = connection.execute(
                f"SELECT id, name FROM {table_name} ORDER BY name"
            ).fetchall()

        return [dict(row) for row in rows]

    def get_pickup_points(self) -> list[dict[str, Any]]:
        with get_connection() as connection:
            rows = connection.execute(
                "SELECT id, address FROM pickup_points ORDER BY id"
            ).fetchall()

        return [dict(row) for row in rows]

    def get_statuses(self) -> list[dict[str, Any]]:
        with get_connection() as connection:
            rows = connection.execute(
                "SELECT id, name FROM order_statuses ORDER BY id"
            ).fetchall()

        return [dict(row) for row in rows]


class ProductRepository:
    SORT_FIELDS = {
        "Без сортировки": "products.id",
        "Количество на складе": "products.stock_quantity",
        "Цена": "products.price",
        "Действующая скидка": "products.discount",
    }

    def list_products(
        self,
        search_text: str = "",
        discount_range_name: str = "Все диапазоны",
        sort_name: str = "Без сортировки",
        sort_direction: str = "ASC",
    ) -> list[dict[str, Any]]:
        query = [
            """
            SELECT
                products.id,
                products.article,
                products.name,
                products.price,
                products.discount,
                products.stock_quantity,
                products.description,
                products.photo_path,
                units.name AS unit_name,
                suppliers.name AS supplier_name,
                manufacturers.name AS manufacturer_name,
                categories.name AS category_name
            FROM products
            JOIN units ON units.id = products.unit_id
            JOIN suppliers ON suppliers.id = products.supplier_id
            JOIN manufacturers ON manufacturers.id = products.manufacturer_id
            JOIN categories ON categories.id = products.category_id
            """
        ]
        conditions = []
        params: list[Any] = []

        discount_range = DISCOUNT_RANGES.get(discount_range_name)
        if discount_range:
            conditions.append("products.discount BETWEEN ? AND ?")
            params.extend(discount_range)

        if conditions:
            query.append("WHERE " + " AND ".join(conditions))

        sort_column = self.SORT_FIELDS.get(sort_name, "products.id")
        direction = "DESC" if sort_direction == "DESC" else "ASC"
        query.append(f"ORDER BY {sort_column} {direction}, products.id ASC")

        with get_connection() as connection:
            rows = connection.execute("\n".join(query), params).fetchall()

        products = [dict(row) for row in rows]
        for product in products:
            product["final_price"] = round(
                product["price"] * (1 - product["discount"] / 100),
                2,
            )

        normalized_search = search_text.strip().casefold()
        if normalized_search:
            tokens = normalized_search.split()
            products = [
                product
                for product in products
                if self._matches_search(product, tokens)
            ]

        return products

    def get_product(self, product_id: int) -> dict[str, Any] | None:
        products = self.list_products()
        for product in products:
            if product["id"] == product_id:
                return product

        return None

    def save_product(self, data: dict[str, Any]) -> int:
        self._validate_product(data)

        with get_connection() as connection:
            unit_id = self._get_or_create(connection, "units", data["unit_name"])
            supplier_id = self._get_or_create(connection, "suppliers", data["supplier_name"])
            manufacturer_id = self._get_or_create(connection, "manufacturers", data["manufacturer_name"])
            category_id = self._get_or_create(connection, "categories", data["category_name"])

            if data.get("id"):
                connection.execute(
                    """
                    UPDATE products
                    SET article = ?, name = ?, unit_id = ?, price = ?,
                        supplier_id = ?, manufacturer_id = ?, category_id = ?,
                        discount = ?, stock_quantity = ?, description = ?, photo_path = ?
                    WHERE id = ?
                    """,
                    (
                        data["article"],
                        data["name"],
                        unit_id,
                        data["price"],
                        supplier_id,
                        manufacturer_id,
                        category_id,
                        data["discount"],
                        data["stock_quantity"],
                        data["description"],
                        data.get("photo_path"),
                        data["id"],
                    ),
                )
                return int(data["id"])

            next_id = connection.execute(
                "SELECT COALESCE(MAX(id), 0) + 1 FROM products"
            ).fetchone()[0]
            cursor = connection.execute(
                """
                INSERT INTO products(
                    id, article, name, unit_id, price, supplier_id, manufacturer_id,
                    category_id, discount, stock_quantity, description, photo_path
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    next_id,
                    data["article"],
                    data["name"],
                    unit_id,
                    data["price"],
                    supplier_id,
                    manufacturer_id,
                    category_id,
                    data["discount"],
                    data["stock_quantity"],
                    data["description"],
                    data.get("photo_path"),
                ),
            )
            return cursor.lastrowid

    def delete_product(self, product_id: int) -> None:
        with get_connection() as connection:
            usage_count = connection.execute(
                "SELECT COUNT(*) FROM order_items WHERE product_id = ?",
                (product_id,),
            ).fetchone()[0]

            if usage_count > 0:
                raise ValueError("Товар присутствует в заказе, поэтому удалить его нельзя.")

            connection.execute("DELETE FROM products WHERE id = ?", (product_id,))

    def article_exists(self, article: str, product_id: int | None = None) -> bool:
        query = "SELECT id FROM products WHERE article = ?"
        params: list[Any] = [article]

        if product_id:
            query += " AND id <> ?"
            params.append(product_id)

        with get_connection() as connection:
            row = connection.execute(query, params).fetchone()

        return row is not None

    def _validate_product(self, data: dict[str, Any]) -> None:
        required_fields = {
            "article": "артикул",
            "name": "наименование",
            "unit_name": "единица измерения",
            "supplier_name": "поставщик",
            "manufacturer_name": "производитель",
            "category_name": "категория",
        }

        for key, label in required_fields.items():
            if not str(data.get(key, "")).strip():
                raise ValueError(f"Заполните поле: {label}.")

        if data["price"] < 0:
            raise ValueError("Стоимость товара не может быть отрицательной.")

        if data["stock_quantity"] < 0:
            raise ValueError("Количество на складе не может быть отрицательным.")

        if data["discount"] < 0 or data["discount"] > 100:
            raise ValueError("Скидка должна быть в диапазоне от 0 до 100%.")

        if self.article_exists(data["article"], data.get("id")):
            raise ValueError("Товар с таким артикулом уже существует.")

    def _get_or_create(self, connection, table_name: str, name: str) -> int:
        name = str(name).strip()
        row = connection.execute(f"SELECT id FROM {table_name} WHERE name = ?", (name,)).fetchone()
        if row:
            return row["id"]

        cursor = connection.execute(f"INSERT INTO {table_name}(name) VALUES (?)", (name,))
        return cursor.lastrowid

    def _matches_search(self, product: dict[str, Any], tokens: list[str]) -> bool:
        # Python casefold корректно работает с кириллицей, в отличие от SQLite LOWER.
        searchable_text = " ".join(
            [
                product["article"],
                product["name"],
                product["description"],
                product["unit_name"],
                product["supplier_name"],
                product["manufacturer_name"],
                product["category_name"],
            ]
        ).casefold()

        return all(token in searchable_text for token in tokens)


class OrderRepository:
    def list_orders(self) -> list[dict[str, Any]]:
        if DB_ENGINE == "mysql":
            items_expression = "GROUP_CONCAT(CONCAT(products.article, ' x ', order_items.quantity) SEPARATOR ', ')"
        else:
            items_expression = "GROUP_CONCAT(products.article || ' x ' || order_items.quantity, ', ')"

        with get_connection() as connection:
            rows = connection.execute(
                f"""
                SELECT
                    orders.id,
                    orders.client_name,
                    orders.receive_code,
                    orders.order_date,
                    orders.delivery_date,
                    pickup_points.address AS pickup_address,
                    order_statuses.name AS status_name,
                    {items_expression} AS items_summary
                FROM orders
                JOIN pickup_points ON pickup_points.id = orders.pickup_point_id
                JOIN order_statuses ON order_statuses.id = orders.status_id
                LEFT JOIN order_items ON order_items.order_id = orders.id
                LEFT JOIN products ON products.id = order_items.product_id
                GROUP BY orders.id
                ORDER BY orders.id
                """
            ).fetchall()

        return [dict(row) for row in rows]

    def get_order(self, order_id: int) -> dict[str, Any] | None:
        with get_connection() as connection:
            order = connection.execute(
                """
                SELECT
                    orders.id,
                    orders.pickup_point_id,
                    orders.client_name,
                    orders.receive_code,
                    orders.status_id,
                    orders.order_date,
                    orders.delivery_date
                FROM orders
                WHERE orders.id = ?
                """,
                (order_id,),
            ).fetchone()

            if not order:
                return None

            items = connection.execute(
                """
                SELECT products.article, products.name, order_items.quantity
                FROM order_items
                JOIN products ON products.id = order_items.product_id
                WHERE order_items.order_id = ?
                ORDER BY products.article
                """,
                (order_id,),
            ).fetchall()

        result = dict(order)
        result["items"] = [dict(item) for item in items]
        return result

    def save_order(self, data: dict[str, Any]) -> int:
        self._validate_order(data)

        with get_connection() as connection:
            if data.get("id"):
                order_id = int(data["id"])
                connection.execute(
                    """
                    UPDATE orders
                    SET pickup_point_id = ?, client_name = ?, receive_code = ?,
                        status_id = ?, order_date = ?, delivery_date = ?
                    WHERE id = ?
                    """,
                    (
                        data["pickup_point_id"],
                        data["client_name"],
                        data["receive_code"],
                        data["status_id"],
                        data["order_date"],
                        data["delivery_date"],
                        order_id,
                    ),
                )
                connection.execute("DELETE FROM order_items WHERE order_id = ?", (order_id,))
            else:
                cursor = connection.execute(
                    """
                    INSERT INTO orders(
                        pickup_point_id, client_name, receive_code,
                        status_id, order_date, delivery_date
                    )
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (
                        data["pickup_point_id"],
                        data["client_name"],
                        data["receive_code"],
                        data["status_id"],
                        data["order_date"],
                        data["delivery_date"],
                    ),
                )
                order_id = cursor.lastrowid

            for item in data["items"]:
                product_id = self._get_product_id(connection, item["article"])
                connection.execute(
                    """
                    INSERT INTO order_items(order_id, product_id, quantity)
                    VALUES (?, ?, ?)
                    """,
                    (order_id, product_id, item["quantity"]),
                )

            return order_id

    def delete_order(self, order_id: int) -> None:
        with get_connection() as connection:
            connection.execute("DELETE FROM orders WHERE id = ?", (order_id,))

    def parse_items_text(self, text: str) -> list[dict[str, Any]]:
        normalized = text.replace(";", ",").replace("\n", ",")
        parts = [part.strip() for part in normalized.split(",") if part.strip()]
        if not parts:
            return []

        if len(parts) == 1:
            return [{"article": parts[0], "quantity": 1}]

        if len(parts) % 2 != 0:
            raise ValueError(
                "Состав заказа укажите как один артикул или парами: артикул, количество."
            )

        items = []
        for index in range(0, len(parts), 2):
            try:
                quantity = int(float(parts[index + 1].replace(",", ".")))
            except ValueError as error:
                raise ValueError("Количество в составе заказа должно быть числом.") from error

            items.append({"article": parts[index], "quantity": quantity})

        return items

    def _validate_order(self, data: dict[str, Any]) -> None:
        if not data.get("pickup_point_id"):
            raise ValueError("Выберите пункт выдачи.")

        if not data.get("status_id"):
            raise ValueError("Выберите статус заказа.")

        if not data.get("items"):
            raise ValueError("Добавьте хотя бы одну позицию заказа.")

        articles = set()
        for item in data["items"]:
            article = str(item.get("article", "")).strip()
            quantity = int(item.get("quantity", 0))

            if not article:
                raise ValueError("У каждой позиции должен быть артикул товара.")

            if quantity <= 0:
                raise ValueError("Количество товара в заказе должно быть больше нуля.")

            if article in articles:
                raise ValueError("Один артикул не должен повторяться в заказе.")

            articles.add(article)

            with get_connection() as connection:
                if self._get_product_id(connection, article) is None:
                    raise ValueError(f"Товар с артикулом {article} не найден.")

    def _get_product_id(self, connection, article: str) -> int | None:
        row = connection.execute(
            "SELECT id FROM products WHERE article = ?",
            (article,),
        ).fetchone()

        return row["id"] if row else None


def resolve_asset_path(relative_path: str | None) -> Path:
    if not relative_path:
        return ASSETS_DIR / "picture.png"

    path = ASSETS_DIR.parent / relative_path
    if path.exists():
        return path

    return ASSETS_DIR / "picture.png"
''',
    # ============================================================
    # END FILE: app/repositories.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: app/ui.py
    # ============================================================
    'app/ui.py': r'''from __future__ import annotations

import os
import re
import shutil
import time
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, font, messagebox, ttk

try:
    from PIL import Image, ImageDraw, ImageTk

    PIL_AVAILABLE = True
except ModuleNotFoundError:
    Image = None
    ImageDraw = None
    ImageTk = None
    PIL_AVAILABLE = False

from config import (
    APP_TITLE,
    ASSETS_DIR,
    COLOR_ACCENT,
    COLOR_DISCOUNT,
    COLOR_OUT_OF_STOCK,
    COLOR_PANEL,
    COLOR_SECONDARY,
    COLOR_WHITE,
    DISCOUNT_RANGES,
    FONT_FAMILY,
    ICON_PATH,
    LOGO_PATH,
    PLACEHOLDER_PATH,
    PRODUCT_IMAGES_DIR,
    ROLE_ADMIN,
    ROLE_CLIENT,
    ROLE_GUEST,
    ROLE_MANAGER,
)
from repositories import (
    AuthRepository,
    CurrentUser,
    LookupRepository,
    OrderRepository,
    ProductRepository,
    resolve_asset_path,
)


def ensure_placeholder_image() -> Path | None:
    if not PIL_AVAILABLE:
        return None

    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    if PLACEHOLDER_PATH.exists():
        return PLACEHOLDER_PATH

    image = Image.new("RGB", (400, 300), "#EFFFFF")
    draw = ImageDraw.Draw(image)
    draw.rectangle((8, 8, 392, 292), outline="#00FFFF", width=4)
    draw.text((118, 132), "Нет изображения", fill="#0000FF")
    image.save(PLACEHOLDER_PATH)
    return PLACEHOLDER_PATH


def load_photo(path: Path, size: tuple[int, int]):
    if not PIL_AVAILABLE:
        return None

    source_path = path if path.exists() else ensure_placeholder_image()
    if not source_path or not source_path.exists():
        return None

    image = Image.open(source_path)
    image.thumbnail(size)
    canvas = Image.new("RGBA", size, (255, 255, 255, 0))
    x = (size[0] - image.width) // 2
    y = (size[1] - image.height) // 2
    canvas.paste(image.convert("RGBA"), (x, y))
    return ImageTk.PhotoImage(canvas)


def show_error(title: str, error: Exception | str) -> None:
    messagebox.showerror(title, str(error))


def show_info(title: str, text: str) -> None:
    messagebox.showinfo(title, text)


def ask_confirmation(title: str, text: str) -> bool:
    return messagebox.askyesno(title, text, icon="warning")


def normalize_date(value: str) -> str | None:
    text = value.strip()
    if not text:
        return None

    for date_format in ("%Y-%m-%d", "%d.%m.%Y"):
        try:
            return datetime.strptime(text, date_format).strftime("%Y-%m-%d")
        except ValueError:
            continue

    raise ValueError("Дата должна быть в формате ГГГГ-ММ-ДД или ДД.ММ.ГГГГ.")


class Application:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.current_frame: tk.Frame | None = None
        self.current_user: CurrentUser | None = None
        self.product_form_window: ProductFormWindow | None = None
        self.order_form_window: OrderFormWindow | None = None

        self.root.geometry("1180x760")
        self.root.minsize(980, 650)
        self.root.configure(bg=COLOR_WHITE)

        if ICON_PATH.exists():
            self.root.iconbitmap(str(ICON_PATH))

        style = ttk.Style()
        style.theme_use("clam")
        style.configure(".", font=(FONT_FAMILY, 10))
        style.configure(
            "Accent.TButton",
            font=(FONT_FAMILY, 10, "bold"),
            background=COLOR_ACCENT,
            foreground=COLOR_WHITE,
        )
        style.map(
            "Accent.TButton",
            background=[("active", COLOR_ACCENT)],
            foreground=[("active", COLOR_WHITE)],
        )

    def show_login(self) -> None:
        self.current_user = None
        self.root.title(f"{APP_TITLE} - вход")
        self._show_frame(LoginFrame(self.root, self))

    def show_products(self, user: CurrentUser) -> None:
        self.current_user = user
        self.root.title(f"{APP_TITLE} - товары")
        self._show_frame(ProductListFrame(self.root, self, user))

    def show_orders(self) -> None:
        if not self.current_user:
            self.show_login()
            return

        self.root.title(f"{APP_TITLE} - заказы")
        self._show_frame(OrderListFrame(self.root, self, self.current_user))

    def logout(self) -> None:
        self.show_login()

    def open_product_form(self, product_id: int | None, refresh_callback) -> None:
        if self.product_form_window and self.product_form_window.winfo_exists():
            self.product_form_window.focus()
            show_info(
                "Редактирование товара",
                "Уже открыто окно добавления или редактирования товара. Закройте его перед продолжением.",
            )
            return

        self.product_form_window = ProductFormWindow(
            self.root,
            product_id,
            refresh_callback,
            self._clear_product_form_reference,
        )

    def open_order_form(self, order_id: int | None, refresh_callback) -> None:
        if self.order_form_window and self.order_form_window.winfo_exists():
            self.order_form_window.focus()
            show_info(
                "Редактирование заказа",
                "Уже открыто окно добавления или редактирования заказа. Закройте его перед продолжением.",
            )
            return

        self.order_form_window = OrderFormWindow(
            self.root,
            order_id,
            refresh_callback,
            self._clear_order_form_reference,
        )

    def _clear_product_form_reference(self) -> None:
        self.product_form_window = None

    def _clear_order_form_reference(self) -> None:
        self.order_form_window = None

    def _show_frame(self, frame: tk.Frame) -> None:
        if self.current_frame:
            self.current_frame.destroy()

        self.current_frame = frame
        self.current_frame.pack(fill="both", expand=True)


class LoginFrame(tk.Frame):
    def __init__(self, master: tk.Misc, app: Application) -> None:
        super().__init__(master, bg=COLOR_WHITE)
        self.app = app
        self.auth_repository = AuthRepository()
        self.login_var = tk.StringVar()
        self.password_var = tk.StringVar()

        self._build()

    def _build(self) -> None:
        container = tk.Frame(self, bg=COLOR_WHITE)
        container.place(relx=0.5, rely=0.5, anchor="center")

        if LOGO_PATH.exists():
            self.logo_image = load_photo(LOGO_PATH, (260, 155))
            if self.logo_image:
                tk.Label(container, image=self.logo_image, bg=COLOR_WHITE).pack(pady=(0, 18))

        tk.Label(
            container,
            text="Вход в систему",
            font=(FONT_FAMILY, 22, "bold"),
            bg=COLOR_WHITE,
            fg=COLOR_ACCENT,
        ).pack(pady=(0, 18))

        form = tk.Frame(container, bg=COLOR_WHITE)
        form.pack()

        tk.Label(form, text="Логин", font=(FONT_FAMILY, 11), bg=COLOR_WHITE).grid(row=0, column=0, sticky="w")
        login_entry = ttk.Entry(form, textvariable=self.login_var, width=38)
        login_entry.grid(row=1, column=0, pady=(4, 12))

        tk.Label(form, text="Пароль", font=(FONT_FAMILY, 11), bg=COLOR_WHITE).grid(row=2, column=0, sticky="w")
        password_entry = ttk.Entry(form, textvariable=self.password_var, show="*", width=38)
        password_entry.grid(row=3, column=0, pady=(4, 18))

        ttk.Button(form, text="Войти", command=self._login, style="Accent.TButton").grid(row=4, column=0, sticky="ew")
        ttk.Button(form, text="Войти как гость", command=self._guest_login).grid(row=5, column=0, sticky="ew", pady=(8, 0))

        login_entry.focus()
        password_entry.bind("<Return>", lambda _event: self._login())

    def _login(self) -> None:
        if not self.login_var.get().strip() or not self.password_var.get().strip():
            show_error(
                "Ошибка входа",
                "Введите логин и пароль. Эти данные берутся из базы данных пользователей.",
            )
            return

        user = self.auth_repository.authenticate(self.login_var.get(), self.password_var.get())
        if not user:
            show_error(
                "Ошибка входа",
                "Пользователь с таким логином и паролем не найден. Проверьте введенные данные.",
            )
            return

        self.app.show_products(user)

    def _guest_login(self) -> None:
        self.app.show_products(CurrentUser(None, "Гость", ROLE_GUEST))


class HeaderFrame(tk.Frame):
    def __init__(self, master: tk.Misc, title: str, user: CurrentUser, app: Application) -> None:
        super().__init__(master, bg=COLOR_SECONDARY)

        tk.Label(
            self,
            text=title,
            font=(FONT_FAMILY, 16, "bold"),
            bg=COLOR_SECONDARY,
            fg="black",
        ).pack(side="left", padx=14, pady=10)

        right = tk.Frame(self, bg=COLOR_SECONDARY)
        right.pack(side="right", padx=14)

        tk.Label(
            right,
            text=user.full_name,
            font=(FONT_FAMILY, 10, "bold"),
            bg=COLOR_SECONDARY,
            fg="black",
        ).pack(side="left", padx=(0, 12))

        ttk.Button(right, text="Выйти", command=app.logout).pack(side="left")


class ProductListFrame(tk.Frame):
    def __init__(self, master: tk.Misc, app: Application, user: CurrentUser) -> None:
        super().__init__(master, bg=COLOR_WHITE)
        self.app = app
        self.user = user
        self.repository = ProductRepository()
        self.search_var = tk.StringVar()
        self.discount_var = tk.StringVar(value="Все диапазоны")
        self.sort_var = tk.StringVar(value="Без сортировки")
        self.direction_var = tk.StringVar(value="ASC")
        self.image_refs: list = []

        self.can_search = user.role_name in (ROLE_MANAGER, ROLE_ADMIN)
        self.can_admin = user.role_name == ROLE_ADMIN

        self._build()
        self.refresh_products()

    def _build(self) -> None:
        HeaderFrame(self, "Список товаров", self.user, self.app).pack(fill="x")

        toolbar = tk.Frame(self, bg=COLOR_PANEL)
        toolbar.pack(fill="x", padx=12, pady=10)

        if self.can_search:
            tk.Label(toolbar, text="Поиск", bg=COLOR_PANEL, font=(FONT_FAMILY, 10, "bold")).grid(row=0, column=0, sticky="w", padx=8, pady=(8, 2))
            search_entry = ttk.Entry(toolbar, textvariable=self.search_var, width=34)
            search_entry.grid(row=1, column=0, padx=8, pady=(0, 8), sticky="ew")

            tk.Label(toolbar, text="Скидка", bg=COLOR_PANEL, font=(FONT_FAMILY, 10, "bold")).grid(row=0, column=1, sticky="w", padx=8, pady=(8, 2))
            discount_combo = ttk.Combobox(
                toolbar,
                textvariable=self.discount_var,
                values=list(DISCOUNT_RANGES.keys()),
                state="readonly",
                width=18,
            )
            discount_combo.grid(row=1, column=1, padx=8, pady=(0, 8))

            tk.Label(toolbar, text="Сортировка", bg=COLOR_PANEL, font=(FONT_FAMILY, 10, "bold")).grid(row=0, column=2, sticky="w", padx=8, pady=(8, 2))
            sort_combo = ttk.Combobox(
                toolbar,
                textvariable=self.sort_var,
                values=list(ProductRepository.SORT_FIELDS.keys()),
                state="readonly",
                width=24,
            )
            sort_combo.grid(row=1, column=2, padx=8, pady=(0, 8))

            tk.Label(toolbar, text="Порядок", bg=COLOR_PANEL, font=(FONT_FAMILY, 10, "bold")).grid(row=0, column=3, sticky="w", padx=8, pady=(8, 2))
            direction_combo = ttk.Combobox(
                toolbar,
                textvariable=self.direction_var,
                values=["ASC", "DESC"],
                state="readonly",
                width=10,
            )
            direction_combo.grid(row=1, column=3, padx=8, pady=(0, 8))

            for variable in (self.search_var, self.discount_var, self.sort_var, self.direction_var):
                variable.trace_add("write", lambda *_args: self.refresh_products())

        if self.can_admin:
            ttk.Button(
                toolbar,
                text="Добавить товар",
                command=lambda: self.app.open_product_form(None, self.refresh_products),
                style="Accent.TButton",
            ).grid(row=1, column=4, padx=8, pady=(0, 8))

        if self.user.role_name in (ROLE_MANAGER, ROLE_ADMIN):
            ttk.Button(toolbar, text="Заказы", command=self.app.show_orders).grid(row=1, column=5, padx=8, pady=(0, 8))

        toolbar.columnconfigure(0, weight=1)

        self.canvas = tk.Canvas(self, bg=COLOR_WHITE, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.products_frame = tk.Frame(self.canvas, bg=COLOR_WHITE)
        self.products_frame.bind(
            "<Configure>",
            lambda _event: self.canvas.configure(scrollregion=self.canvas.bbox("all")),
        )
        self.canvas_window = self.canvas.create_window((0, 0), window=self.products_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True, padx=(12, 0), pady=(0, 12))
        self.scrollbar.pack(side="right", fill="y", padx=(0, 12), pady=(0, 12))
        self.canvas.bind("<Configure>", self._resize_canvas_window)
        self.canvas.bind_all("<MouseWheel>", self._on_mouse_wheel)

    def refresh_products(self) -> None:
        for child in self.products_frame.winfo_children():
            child.destroy()

        self.image_refs.clear()

        try:
            products = self.repository.list_products(
                self.search_var.get() if self.can_search else "",
                self.discount_var.get() if self.can_search else "Все диапазоны",
                self.sort_var.get() if self.can_search else "Без сортировки",
                self.direction_var.get() if self.can_search else "ASC",
            )
        except Exception as error:
            show_error("Ошибка загрузки товаров", error)
            return

        if not products:
            tk.Label(
                self.products_frame,
                text="Товары не найдены",
                font=(FONT_FAMILY, 13),
                bg=COLOR_WHITE,
            ).pack(pady=30)
            return

        for product in products:
            self._create_product_card(product)

    def _create_product_card(self, product: dict) -> None:
        background = COLOR_WHITE
        foreground = "black"

        if product["stock_quantity"] == 0:
            background = COLOR_OUT_OF_STOCK
        elif product["discount"] > 15:
            background = COLOR_DISCOUNT
            foreground = COLOR_WHITE

        card = tk.Frame(
            self.products_frame,
            bg=background,
            bd=1,
            relief="solid",
            padx=10,
            pady=10,
        )
        card.pack(fill="x", pady=6)
        card.columnconfigure(1, weight=1)

        image_path = resolve_asset_path(product.get("photo_path"))
        photo = load_photo(image_path, (120, 86))
        if photo:
            self.image_refs.append(photo)
            tk.Label(card, image=photo, bg=background).grid(row=0, column=0, rowspan=5, sticky="nw", padx=(0, 12))
        else:
            tk.Label(
                card,
                text="Нет\nизображения",
                width=16,
                height=5,
                bg="#EEEEEE",
                fg=COLOR_ACCENT,
                font=(FONT_FAMILY, 9, "bold"),
            ).grid(row=0, column=0, rowspan=5, sticky="nw", padx=(0, 12))

        title = f"{product['article']} - {product['name']}"
        tk.Label(
            card,
            text=title,
            font=(FONT_FAMILY, 12, "bold"),
            bg=background,
            fg=foreground,
            anchor="w",
            justify="left",
            wraplength=780,
        ).grid(row=0, column=1, sticky="ew")

        details = (
            f"Категория: {product['category_name']} | Производитель: {product['manufacturer_name']} | "
            f"Поставщик: {product['supplier_name']} | Ед.: {product['unit_name']}"
        )
        tk.Label(
            card,
            text=details,
            bg=background,
            fg=foreground,
            anchor="w",
            justify="left",
            wraplength=780,
        ).grid(row=1, column=1, sticky="ew", pady=(4, 0))

        tk.Label(
            card,
            text=product["description"],
            bg=background,
            fg=foreground,
            anchor="w",
            justify="left",
            wraplength=780,
        ).grid(row=2, column=1, sticky="ew", pady=(4, 0))

        price_frame = tk.Frame(card, bg=background)
        price_frame.grid(row=3, column=1, sticky="w", pady=(8, 0))

        if product["discount"] > 0:
            old_font = font.Font(family=FONT_FAMILY, size=10, overstrike=True)
            tk.Label(
                price_frame,
                text=f"{product['price']:.2f} руб.",
                font=old_font,
                fg="red",
                bg=background,
            ).pack(side="left")
            tk.Label(
                price_frame,
                text=f"  {product['final_price']:.2f} руб.",
                font=(FONT_FAMILY, 10, "bold"),
                fg="black" if background != COLOR_OUT_OF_STOCK else "black",
                bg=background,
            ).pack(side="left")
        else:
            tk.Label(
                price_frame,
                text=f"{product['price']:.2f} руб.",
                font=(FONT_FAMILY, 10, "bold"),
                bg=background,
                fg=foreground,
            ).pack(side="left")

        bottom = f"На складе: {product['stock_quantity']} | Действующая скидка: {product['discount']:.2f}%"
        tk.Label(card, text=bottom, bg=background, fg=foreground, anchor="w").grid(row=4, column=1, sticky="w")

        actions = tk.Frame(card, bg=background)
        actions.grid(row=0, column=2, rowspan=5, sticky="ne", padx=(10, 0))

        if self.can_admin:
            ttk.Button(
                actions,
                text="Редактировать",
                command=lambda product_id=product["id"]: self.app.open_product_form(product_id, self.refresh_products),
            ).pack(fill="x", pady=(0, 6))
            ttk.Button(
                actions,
                text="Удалить",
                command=lambda product_id=product["id"]: self._delete_product(product_id),
            ).pack(fill="x")

            self._bind_card_edit(card, product["id"])

    def _bind_card_edit(self, widget: tk.Widget, product_id: int) -> None:
        if isinstance(widget, ttk.Button):
            return

        widget.bind(
            "<Double-Button-1>",
            lambda _event: self.app.open_product_form(product_id, self.refresh_products),
        )

        for child in widget.winfo_children():
            self._bind_card_edit(child, product_id)

    def _delete_product(self, product_id: int) -> None:
        if not ask_confirmation(
            "Удаление товара",
            "Удаление товара необратимо. Вы действительно хотите удалить выбранный товар?",
        ):
            return

        try:
            self.repository.delete_product(product_id)
            self.refresh_products()
            show_info("Удаление товара", "Товар удален.")
        except Exception as error:
            show_error("Ошибка удаления товара", error)

    def _resize_canvas_window(self, event) -> None:
        self.canvas.itemconfigure(self.canvas_window, width=event.width)

    def _on_mouse_wheel(self, event) -> None:
        if self.winfo_ismapped():
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")


class ProductFormWindow(tk.Toplevel):
    def __init__(
        self,
        master: tk.Misc,
        product_id: int | None,
        refresh_callback,
        close_callback,
    ) -> None:
        super().__init__(master)
        self.product_id = product_id
        self.refresh_callback = refresh_callback
        self.close_callback = close_callback
        self.repository = ProductRepository()
        self.lookup_repository = LookupRepository()
        self.selected_image_path: Path | None = None
        self.current_photo_path: str | None = None
        self.preview_image = None

        self.title("Редактирование товара" if product_id else "Добавление товара")
        self.geometry("740x720")
        self.minsize(680, 650)
        self.configure(bg=COLOR_WHITE)
        self.transient(master)
        self.grab_set()

        if ICON_PATH.exists():
            self.iconbitmap(str(ICON_PATH))

        self.protocol("WM_DELETE_WINDOW", self._close)
        self._build()
        self._load_product()

    def _build(self) -> None:
        container = tk.Frame(self, bg=COLOR_WHITE, padx=16, pady=16)
        container.pack(fill="both", expand=True)
        container.columnconfigure(1, weight=1)

        title = "Редактирование товара" if self.product_id else "Добавление товара"
        tk.Label(
            container,
            text=title,
            font=(FONT_FAMILY, 16, "bold"),
            fg=COLOR_ACCENT,
            bg=COLOR_WHITE,
        ).grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 14))

        self.id_var = tk.StringVar()
        self.article_var = tk.StringVar()
        self.name_var = tk.StringVar()
        self.category_var = tk.StringVar()
        self.manufacturer_var = tk.StringVar()
        self.supplier_var = tk.StringVar()
        self.price_var = tk.StringVar()
        self.unit_var = tk.StringVar()
        self.stock_var = tk.StringVar()
        self.discount_var = tk.StringVar()

        row = 1
        if self.product_id:
            self._add_labeled_entry(container, row, "ID", self.id_var, readonly=True)
            row += 1

        self._add_labeled_entry(container, row, "Артикул", self.article_var)
        row += 1
        self._add_labeled_entry(container, row, "Наименование товара", self.name_var)
        row += 1

        category_values = [item["name"] for item in self.lookup_repository.get_values("categories")]
        supplier_values = [item["name"] for item in self.lookup_repository.get_values("suppliers")]
        manufacturer_values = [item["name"] for item in self.lookup_repository.get_values("manufacturers")]
        unit_values = [item["name"] for item in self.lookup_repository.get_values("units")]

        self._add_labeled_combo(container, row, "Категория товара", self.category_var, category_values)
        row += 1
        self._add_labeled_combo(container, row, "Производитель", self.manufacturer_var, manufacturer_values)
        row += 1
        self._add_labeled_combo(container, row, "Поставщик", self.supplier_var, supplier_values)
        row += 1
        self._add_labeled_entry(container, row, "Цена", self.price_var)
        row += 1
        self._add_labeled_combo(container, row, "Единица измерения", self.unit_var, unit_values)
        row += 1
        self._add_labeled_entry(container, row, "Количество на складе", self.stock_var)
        row += 1
        self._add_labeled_entry(container, row, "Действующая скидка", self.discount_var)
        row += 1

        tk.Label(container, text="Описание товара", bg=COLOR_WHITE, font=(FONT_FAMILY, 10, "bold")).grid(row=row, column=0, sticky="nw", pady=6)
        self.description_text = tk.Text(container, height=5, wrap="word", font=(FONT_FAMILY, 10))
        self.description_text.grid(row=row, column=1, columnspan=2, sticky="nsew", pady=6)
        row += 1

        tk.Label(container, text="Фото товара", bg=COLOR_WHITE, font=(FONT_FAMILY, 10, "bold")).grid(row=row, column=0, sticky="nw", pady=6)
        photo_frame = tk.Frame(container, bg=COLOR_WHITE)
        photo_frame.grid(row=row, column=1, columnspan=2, sticky="w", pady=6)
        self.photo_label = tk.Label(photo_frame, bg=COLOR_WHITE, width=180, height=120)
        self.photo_label.pack(side="left", padx=(0, 12))
        ttk.Button(photo_frame, text="Выбрать фото", command=self._choose_photo).pack(side="left")
        row += 1

        buttons = tk.Frame(container, bg=COLOR_WHITE)
        buttons.grid(row=row, column=0, columnspan=3, sticky="e", pady=(16, 0))
        ttk.Button(buttons, text="Сохранить", command=self._save, style="Accent.TButton").pack(side="left", padx=(0, 8))
        ttk.Button(buttons, text="Отмена", command=self._close).pack(side="left")

    def _add_labeled_entry(
        self,
        container: tk.Frame,
        row: int,
        label: str,
        variable: tk.StringVar,
        readonly: bool = False,
    ) -> None:
        tk.Label(container, text=label, bg=COLOR_WHITE, font=(FONT_FAMILY, 10, "bold")).grid(row=row, column=0, sticky="w", pady=6)
        state = "readonly" if readonly else "normal"
        entry = ttk.Entry(container, textvariable=variable, state=state)
        entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=6)

    def _add_labeled_combo(
        self,
        container: tk.Frame,
        row: int,
        label: str,
        variable: tk.StringVar,
        values: list[str],
    ) -> None:
        tk.Label(container, text=label, bg=COLOR_WHITE, font=(FONT_FAMILY, 10, "bold")).grid(row=row, column=0, sticky="w", pady=6)
        combo = ttk.Combobox(container, textvariable=variable, values=values, state="readonly")
        combo.grid(row=row, column=1, columnspan=2, sticky="ew", pady=6)
        if values and not variable.get():
            variable.set(values[0])

    def _load_product(self) -> None:
        if not self.product_id:
            self.price_var.set("0")
            self.stock_var.set("0")
            self.discount_var.set("0")
            self._set_photo_preview(None)
            return

        product = self.repository.get_product(self.product_id)
        if not product:
            show_error("Редактирование товара", "Выбранный товар не найден.")
            self._close()
            return

        self.id_var.set(str(product["id"]))
        self.article_var.set(product["article"])
        self.name_var.set(product["name"])
        self.category_var.set(product["category_name"])
        self.manufacturer_var.set(product["manufacturer_name"])
        self.supplier_var.set(product["supplier_name"])
        self.price_var.set(str(product["price"]))
        self.unit_var.set(product["unit_name"])
        self.stock_var.set(str(product["stock_quantity"]))
        self.discount_var.set(str(product["discount"]))
        self.description_text.insert("1.0", product["description"])
        self.current_photo_path = product.get("photo_path")
        self._set_photo_preview(self.current_photo_path)

    def _set_photo_preview(self, relative_path: str | None) -> None:
        path = resolve_asset_path(relative_path)
        self.preview_image = load_photo(path, (180, 120))
        if self.preview_image:
            self.photo_label.configure(image=self.preview_image, text="")
        else:
            self.photo_label.configure(image="", text="Нет изображения", fg=COLOR_ACCENT)

    def _choose_photo(self) -> None:
        file_name = filedialog.askopenfilename(
            title="Выберите изображение товара",
            filetypes=[
                ("Изображения", "*.png *.jpg *.jpeg *.bmp"),
                ("Все файлы", "*.*"),
            ],
        )

        if not file_name:
            return

        self.selected_image_path = Path(file_name)
        self.preview_image = load_photo(self.selected_image_path, (180, 120))
        if self.preview_image:
            self.photo_label.configure(image=self.preview_image, text="")
        else:
            self.photo_label.configure(image="", text=self.selected_image_path.name, fg=COLOR_ACCENT)

    def _save(self) -> None:
        try:
            photo_path = self.current_photo_path
            if self.selected_image_path:
                photo_path = self._save_uploaded_photo(self.selected_image_path, self.current_photo_path)

            data = {
                "id": self.product_id,
                "article": self.article_var.get().strip(),
                "name": self.name_var.get().strip(),
                "category_name": self.category_var.get().strip(),
                "manufacturer_name": self.manufacturer_var.get().strip(),
                "supplier_name": self.supplier_var.get().strip(),
                "price": float(self.price_var.get().replace(",", ".")),
                "unit_name": self.unit_var.get().strip(),
                "stock_quantity": int(float(self.stock_var.get().replace(",", "."))),
                "discount": float(self.discount_var.get().replace(",", ".")),
                "description": self.description_text.get("1.0", "end").strip(),
                "photo_path": photo_path,
            }

            self.repository.save_product(data)
            self.refresh_callback()
            show_info("Сохранение товара", "Данные товара сохранены.")
            self._close()
        except ValueError as error:
            show_error("Ошибка заполнения товара", error)
        except Exception as error:
            show_error("Ошибка сохранения товара", error)

    def _save_uploaded_photo(self, source_path: Path, old_relative_path: str | None) -> str:
        if not PIL_AVAILABLE:
            raise ValueError("Для добавления и сжатия изображений установите пакет Pillow.")

        PRODUCT_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
        safe_article = re.sub(r"[^A-Za-zА-Яа-я0-9_-]+", "_", self.article_var.get().strip()) or "product"
        file_name = f"uploaded_{safe_article}_{int(time.time())}.jpg"
        destination = PRODUCT_IMAGES_DIR / file_name

        with Image.open(source_path) as image:
            image.thumbnail((300, 200))
            background = Image.new("RGB", image.size, (255, 255, 255))
            background.paste(image.convert("RGB"))
            background.save(destination, "JPEG", quality=90)

        if old_relative_path:
            old_path = ASSETS_DIR.parent / old_relative_path
            if old_path.exists() and old_path.parent == PRODUCT_IMAGES_DIR:
                try:
                    old_path.unlink()
                except OSError:
                    pass

        return f"assets/products/{file_name}"

    def _close(self) -> None:
        self.close_callback()
        self.grab_release()
        self.destroy()


class OrderListFrame(tk.Frame):
    def __init__(self, master: tk.Misc, app: Application, user: CurrentUser) -> None:
        super().__init__(master, bg=COLOR_WHITE)
        self.app = app
        self.user = user
        self.repository = OrderRepository()
        self.can_admin = user.role_name == ROLE_ADMIN

        self._build()
        self.refresh_orders()

    def _build(self) -> None:
        HeaderFrame(self, "Заказы", self.user, self.app).pack(fill="x")

        toolbar = tk.Frame(self, bg=COLOR_PANEL)
        toolbar.pack(fill="x", padx=12, pady=10)
        ttk.Button(toolbar, text="Назад к товарам", command=lambda: self.app.show_products(self.user)).pack(side="left", padx=8, pady=8)

        if self.can_admin:
            ttk.Button(
                toolbar,
                text="Добавить заказ",
                command=lambda: self.app.open_order_form(None, self.refresh_orders),
                style="Accent.TButton",
            ).pack(side="left", padx=8, pady=8)
            ttk.Button(toolbar, text="Редактировать", command=self._edit_selected).pack(side="left", padx=8, pady=8)
            ttk.Button(toolbar, text="Удалить", command=self._delete_selected).pack(side="left", padx=8, pady=8)

        columns = ("id", "status", "pickup", "order_date", "delivery_date", "client", "code", "items")
        self.tree = ttk.Treeview(self, columns=columns, show="headings")
        headings = {
            "id": "Номер",
            "status": "Статус",
            "pickup": "Пункт выдачи",
            "order_date": "Дата заказа",
            "delivery_date": "Дата выдачи",
            "client": "Клиент",
            "code": "Код",
            "items": "Артикулы",
        }
        widths = {
            "id": 70,
            "status": 110,
            "pickup": 260,
            "order_date": 110,
            "delivery_date": 110,
            "client": 210,
            "code": 80,
            "items": 260,
        }

        for column in columns:
            self.tree.heading(column, text=headings[column])
            self.tree.column(column, width=widths[column], anchor="w")

        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        self.tree.pack(side="left", fill="both", expand=True, padx=(12, 0), pady=(0, 12))
        scrollbar.pack(side="right", fill="y", padx=(0, 12), pady=(0, 12))

        if self.can_admin:
            self.tree.bind("<Double-Button-1>", lambda _event: self._edit_selected())

    def refresh_orders(self) -> None:
        for item in self.tree.get_children():
            self.tree.delete(item)

        try:
            orders = self.repository.list_orders()
        except Exception as error:
            show_error("Ошибка загрузки заказов", error)
            return

        for order in orders:
            self.tree.insert(
                "",
                "end",
                iid=str(order["id"]),
                values=(
                    order["id"],
                    order["status_name"],
                    order["pickup_address"],
                    order["order_date"] or "",
                    order["delivery_date"] or "",
                    order["client_name"] or "",
                    order["receive_code"] or "",
                    order["items_summary"] or "",
                ),
            )

    def _get_selected_order_id(self) -> int | None:
        selected = self.tree.selection()
        if not selected:
            show_error("Заказы", "Выберите заказ в списке.")
            return None

        return int(selected[0])

    def _edit_selected(self) -> None:
        order_id = self._get_selected_order_id()
        if order_id:
            self.app.open_order_form(order_id, self.refresh_orders)

    def _delete_selected(self) -> None:
        order_id = self._get_selected_order_id()
        if not order_id:
            return

        if not ask_confirmation(
            "Удаление заказа",
            "Удаление заказа необратимо. Вы действительно хотите удалить выбранный заказ?",
        ):
            return

        try:
            self.repository.delete_order(order_id)
            self.refresh_orders()
            show_info("Удаление заказа", "Заказ удален.")
        except Exception as error:
            show_error("Ошибка удаления заказа", error)


class OrderFormWindow(tk.Toplevel):
    def __init__(
        self,
        master: tk.Misc,
        order_id: int | None,
        refresh_callback,
        close_callback,
    ) -> None:
        super().__init__(master)
        self.order_id = order_id
        self.refresh_callback = refresh_callback
        self.close_callback = close_callback
        self.repository = OrderRepository()
        self.lookup_repository = LookupRepository()
        self.product_repository = ProductRepository()
        self.pickup_points = self.lookup_repository.get_pickup_points()
        self.statuses = self.lookup_repository.get_statuses()
        self.products = self.product_repository.list_products()
        self.items: list[dict] = []

        self.title("Редактирование заказа" if order_id else "Добавление заказа")
        self.geometry("820x680")
        self.minsize(760, 620)
        self.configure(bg=COLOR_WHITE)
        self.transient(master)
        self.grab_set()

        if ICON_PATH.exists():
            self.iconbitmap(str(ICON_PATH))

        self.protocol("WM_DELETE_WINDOW", self._close)
        self._build()
        self._load_order()

    def _build(self) -> None:
        container = tk.Frame(self, bg=COLOR_WHITE, padx=16, pady=16)
        container.pack(fill="both", expand=True)
        container.columnconfigure(1, weight=1)

        title = "Редактирование заказа" if self.order_id else "Добавление заказа"
        tk.Label(
            container,
            text=title,
            font=(FONT_FAMILY, 16, "bold"),
            fg=COLOR_ACCENT,
            bg=COLOR_WHITE,
        ).grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 14))

        self.id_var = tk.StringVar()
        self.status_var = tk.StringVar()
        self.pickup_var = tk.StringVar()
        self.client_var = tk.StringVar()
        self.code_var = tk.StringVar()
        self.order_date_var = tk.StringVar()
        self.delivery_date_var = tk.StringVar()
        self.article_var = tk.StringVar()
        self.quantity_var = tk.StringVar(value="1")

        row = 1
        if self.order_id:
            self._add_labeled_entry(container, row, "Номер заказа", self.id_var, readonly=True)
            row += 1

        status_values = [status["name"] for status in self.statuses]
        pickup_values = [f"{point['id']} - {point['address']}" for point in self.pickup_points]
        product_values = [product["article"] for product in self.products]

        self._add_labeled_combo(container, row, "Статус заказа", self.status_var, status_values)
        row += 1
        self._add_labeled_combo(container, row, "Адрес пункта выдачи", self.pickup_var, pickup_values)
        row += 1
        self._add_labeled_entry(container, row, "ФИО клиента", self.client_var)
        row += 1
        self._add_labeled_entry(container, row, "Код для получения", self.code_var)
        row += 1
        self._add_labeled_entry(container, row, "Дата заказа", self.order_date_var)
        row += 1
        self._add_labeled_entry(container, row, "Дата выдачи", self.delivery_date_var)
        row += 1

        items_frame = tk.LabelFrame(container, text="Позиции заказа", bg=COLOR_WHITE, font=(FONT_FAMILY, 10, "bold"))
        items_frame.grid(row=row, column=0, columnspan=3, sticky="nsew", pady=(12, 0))
        items_frame.columnconfigure(0, weight=1)
        row += 1

        editor = tk.Frame(items_frame, bg=COLOR_WHITE)
        editor.grid(row=0, column=0, sticky="ew", padx=8, pady=8)

        ttk.Combobox(editor, textvariable=self.article_var, values=product_values, state="readonly", width=20).pack(side="left")
        ttk.Entry(editor, textvariable=self.quantity_var, width=10).pack(side="left", padx=8)
        ttk.Button(editor, text="Добавить позицию", command=self._add_item).pack(side="left", padx=(0, 8))
        ttk.Button(editor, text="Удалить позицию", command=self._remove_selected_item).pack(side="left")

        item_columns = ("article", "name", "quantity")
        self.items_tree = ttk.Treeview(items_frame, columns=item_columns, show="headings", height=8)
        self.items_tree.heading("article", text="Артикул")
        self.items_tree.heading("name", text="Товар")
        self.items_tree.heading("quantity", text="Количество")
        self.items_tree.column("article", width=140)
        self.items_tree.column("name", width=480)
        self.items_tree.column("quantity", width=110)
        self.items_tree.grid(row=1, column=0, sticky="nsew", padx=8, pady=(0, 8))

        if product_values:
            self.article_var.set(product_values[0])

        buttons = tk.Frame(container, bg=COLOR_WHITE)
        buttons.grid(row=row, column=0, columnspan=3, sticky="e", pady=(16, 0))
        ttk.Button(buttons, text="Сохранить", command=self._save, style="Accent.TButton").pack(side="left", padx=(0, 8))
        ttk.Button(buttons, text="Отмена", command=self._close).pack(side="left")

    def _add_labeled_entry(
        self,
        container: tk.Frame,
        row: int,
        label: str,
        variable: tk.StringVar,
        readonly: bool = False,
    ) -> None:
        tk.Label(container, text=label, bg=COLOR_WHITE, font=(FONT_FAMILY, 10, "bold")).grid(row=row, column=0, sticky="w", pady=6)
        state = "readonly" if readonly else "normal"
        ttk.Entry(container, textvariable=variable, state=state).grid(row=row, column=1, columnspan=2, sticky="ew", pady=6)

    def _add_labeled_combo(
        self,
        container: tk.Frame,
        row: int,
        label: str,
        variable: tk.StringVar,
        values: list[str],
    ) -> None:
        tk.Label(container, text=label, bg=COLOR_WHITE, font=(FONT_FAMILY, 10, "bold")).grid(row=row, column=0, sticky="w", pady=6)
        combo = ttk.Combobox(container, textvariable=variable, values=values, state="readonly")
        combo.grid(row=row, column=1, columnspan=2, sticky="ew", pady=6)
        if values:
            variable.set(values[0])

    def _load_order(self) -> None:
        if not self.order_id:
            self.order_date_var.set(datetime.now().strftime("%Y-%m-%d"))
            return

        order = self.repository.get_order(self.order_id)
        if not order:
            show_error("Редактирование заказа", "Выбранный заказ не найден.")
            self._close()
            return

        self.id_var.set(str(order["id"]))
        self.status_var.set(self._status_name_by_id(order["status_id"]))
        self.pickup_var.set(self._pickup_label_by_id(order["pickup_point_id"]))
        self.client_var.set(order["client_name"] or "")
        self.code_var.set(order["receive_code"] or "")
        self.order_date_var.set(order["order_date"] or "")
        self.delivery_date_var.set(order["delivery_date"] or "")
        self.items = order["items"]
        self._refresh_items_tree()

    def _add_item(self) -> None:
        try:
            article = self.article_var.get().strip()
            quantity = int(float(self.quantity_var.get().replace(",", ".")))

            if quantity <= 0:
                raise ValueError("Количество должно быть больше нуля.")

            if any(item["article"] == article for item in self.items):
                raise ValueError("Этот артикул уже добавлен в заказ.")

            product = next((item for item in self.products if item["article"] == article), None)
            if not product:
                raise ValueError("Выберите артикул товара из списка.")

            self.items.append(
                {
                    "article": article,
                    "name": product["name"],
                    "quantity": quantity,
                }
            )
            self._refresh_items_tree()
        except ValueError as error:
            show_error("Позиция заказа", error)

    def _remove_selected_item(self) -> None:
        selected = self.items_tree.selection()
        if not selected:
            show_error("Позиция заказа", "Выберите позицию для удаления.")
            return

        index = int(selected[0])
        self.items.pop(index)
        self._refresh_items_tree()

    def _refresh_items_tree(self) -> None:
        for item in self.items_tree.get_children():
            self.items_tree.delete(item)

        for index, item in enumerate(self.items):
            self.items_tree.insert(
                "",
                "end",
                iid=str(index),
                values=(item["article"], item.get("name", ""), item["quantity"]),
            )

    def _save(self) -> None:
        try:
            data = {
                "id": self.order_id,
                "pickup_point_id": self._pickup_id_from_label(self.pickup_var.get()),
                "client_name": self.client_var.get().strip(),
                "receive_code": self.code_var.get().strip(),
                "status_id": self._status_id_by_name(self.status_var.get()),
                "order_date": normalize_date(self.order_date_var.get()),
                "delivery_date": normalize_date(self.delivery_date_var.get()),
                "items": [
                    {
                        "article": item["article"],
                        "quantity": int(item["quantity"]),
                    }
                    for item in self.items
                ],
            }
            self.repository.save_order(data)
            self.refresh_callback()
            show_info("Сохранение заказа", "Данные заказа сохранены.")
            self._close()
        except ValueError as error:
            show_error("Ошибка заполнения заказа", error)
        except Exception as error:
            show_error("Ошибка сохранения заказа", error)

    def _status_id_by_name(self, name: str) -> int | None:
        for status in self.statuses:
            if status["name"] == name:
                return status["id"]

        return None

    def _status_name_by_id(self, status_id: int) -> str:
        for status in self.statuses:
            if status["id"] == status_id:
                return status["name"]

        return ""

    def _pickup_id_from_label(self, label: str) -> int | None:
        if " - " not in label:
            return None

        return int(label.split(" - ", 1)[0])

    def _pickup_label_by_id(self, pickup_point_id: int) -> str:
        for point in self.pickup_points:
            if point["id"] == pickup_point_id:
                return f"{point['id']} - {point['address']}"

        return ""

    def _close(self) -> None:
        self.close_callback()
        self.grab_release()
        self.destroy()
''',
    # ============================================================
    # END FILE: app/ui.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: database/schema.sql
    # ============================================================
    'database/schema.sql': r'''PRAGMA foreign_keys = ON;

CREATE TABLE IF NOT EXISTS roles (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    role_id INTEGER NOT NULL,
    full_name TEXT NOT NULL,
    login TEXT NOT NULL UNIQUE,
    password TEXT NOT NULL,
    FOREIGN KEY (role_id) REFERENCES roles(id)
);

CREATE TABLE IF NOT EXISTS categories (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS suppliers (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS manufacturers (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS units (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS products (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    article TEXT NOT NULL UNIQUE,
    name TEXT NOT NULL,
    unit_id INTEGER NOT NULL,
    price REAL NOT NULL CHECK (price >= 0),
    supplier_id INTEGER NOT NULL,
    manufacturer_id INTEGER NOT NULL,
    category_id INTEGER NOT NULL,
    discount REAL NOT NULL DEFAULT 0 CHECK (discount >= 0 AND discount <= 100),
    stock_quantity INTEGER NOT NULL DEFAULT 0 CHECK (stock_quantity >= 0),
    description TEXT NOT NULL DEFAULT '',
    photo_path TEXT,
    FOREIGN KEY (unit_id) REFERENCES units(id),
    FOREIGN KEY (supplier_id) REFERENCES suppliers(id),
    FOREIGN KEY (manufacturer_id) REFERENCES manufacturers(id),
    FOREIGN KEY (category_id) REFERENCES categories(id)
);

CREATE TABLE IF NOT EXISTS pickup_points (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    address TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS order_statuses (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS orders (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    pickup_point_id INTEGER NOT NULL,
    client_user_id INTEGER,
    client_name TEXT,
    receive_code TEXT,
    status_id INTEGER NOT NULL,
    order_date TEXT,
    delivery_date TEXT,
    FOREIGN KEY (pickup_point_id) REFERENCES pickup_points(id),
    FOREIGN KEY (client_user_id) REFERENCES users(id),
    FOREIGN KEY (status_id) REFERENCES order_statuses(id)
);

CREATE TABLE IF NOT EXISTS order_items (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    order_id INTEGER NOT NULL,
    product_id INTEGER NOT NULL,
    quantity INTEGER NOT NULL CHECK (quantity > 0),
    FOREIGN KEY (order_id) REFERENCES orders(id) ON DELETE CASCADE,
    FOREIGN KEY (product_id) REFERENCES products(id),
    UNIQUE (order_id, product_id)
);

CREATE INDEX IF NOT EXISTS idx_products_article ON products(article);
CREATE INDEX IF NOT EXISTS idx_products_name ON products(name);
CREATE INDEX IF NOT EXISTS idx_orders_status ON orders(status_id);
CREATE INDEX IF NOT EXISTS idx_order_items_product ON order_items(product_id);
''',
    # ============================================================
    # END FILE: database/schema.sql
    # ============================================================

    # ============================================================
    # BEGIN FILE: database/schema_sqlite.sql
    # ============================================================
    'database/schema_sqlite.sql': r'''PRAGMA foreign_keys = ON;

CREATE TABLE IF NOT EXISTS roles (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    role_id INTEGER NOT NULL,
    full_name TEXT NOT NULL,
    login TEXT NOT NULL UNIQUE,
    password TEXT NOT NULL,
    FOREIGN KEY (role_id) REFERENCES roles(id)
);

CREATE TABLE IF NOT EXISTS categories (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS suppliers (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS manufacturers (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS units (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS products (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    article TEXT NOT NULL UNIQUE,
    name TEXT NOT NULL,
    unit_id INTEGER NOT NULL,
    price REAL NOT NULL CHECK (price >= 0),
    supplier_id INTEGER NOT NULL,
    manufacturer_id INTEGER NOT NULL,
    category_id INTEGER NOT NULL,
    discount REAL NOT NULL DEFAULT 0 CHECK (discount >= 0 AND discount <= 100),
    stock_quantity INTEGER NOT NULL DEFAULT 0 CHECK (stock_quantity >= 0),
    description TEXT NOT NULL DEFAULT '',
    photo_path TEXT,
    FOREIGN KEY (unit_id) REFERENCES units(id),
    FOREIGN KEY (supplier_id) REFERENCES suppliers(id),
    FOREIGN KEY (manufacturer_id) REFERENCES manufacturers(id),
    FOREIGN KEY (category_id) REFERENCES categories(id)
);

CREATE TABLE IF NOT EXISTS pickup_points (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    address TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS order_statuses (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS orders (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    pickup_point_id INTEGER NOT NULL,
    client_user_id INTEGER,
    client_name TEXT,
    receive_code TEXT,
    status_id INTEGER NOT NULL,
    order_date TEXT,
    delivery_date TEXT,
    FOREIGN KEY (pickup_point_id) REFERENCES pickup_points(id),
    FOREIGN KEY (client_user_id) REFERENCES users(id),
    FOREIGN KEY (status_id) REFERENCES order_statuses(id)
);

CREATE TABLE IF NOT EXISTS order_items (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    order_id INTEGER NOT NULL,
    product_id INTEGER NOT NULL,
    quantity INTEGER NOT NULL CHECK (quantity > 0),
    FOREIGN KEY (order_id) REFERENCES orders(id) ON DELETE CASCADE,
    FOREIGN KEY (product_id) REFERENCES products(id),
    UNIQUE (order_id, product_id)
);

CREATE INDEX IF NOT EXISTS idx_products_article ON products(article);
CREATE INDEX IF NOT EXISTS idx_products_name ON products(name);
CREATE INDEX IF NOT EXISTS idx_orders_status ON orders(status_id);
CREATE INDEX IF NOT EXISTS idx_order_items_product ON order_items(product_id);
''',
    # ============================================================
    # END FILE: database/schema_sqlite.sql
    # ============================================================

    # ============================================================
    # BEGIN FILE: database/seed_sqlite.sql
    # ============================================================
    'database/seed_sqlite.sql': r'''-- Начальные данные для демонстрационного экзамена, вариант 5.
INSERT INTO roles(id, name) VALUES (1, 'Администратор');
INSERT INTO roles(id, name) VALUES (2, 'Менеджер');
INSERT INTO roles(id, name) VALUES (3, 'Авторизированный клиент');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (1, 1, 'Никифорова Анна Семеновна', '94d5ous@gmail.com', 'uzWC67');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (2, 1, 'Стелина Евгения Петровна', 'uth4iz@mail.com', '2L6KZG');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (3, 1, 'Никифорова Весения Николаевна', '5d4zbu@tutanota.com', 'rwVDh9');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (4, 2, 'Сазонов Руслан Германович', 'ptec8ym@yahoo.com', 'LdNyos');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (5, 2, 'Одинцов Серафим Артёмович', '1qz4kw@mail.com', 'gynQMT');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (6, 2, 'Старикова Елена Павловна', '4np6se@mail.com', 'AtnDjr');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (7, 3, 'Степанов Михаил Артёмович', 'yzls62@outlook.com', 'JlFRCZ');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (8, 3, 'Михайлюк Анна Вячеславовна', '1diph5e@tutanota.com', '8ntwUp');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (9, 3, 'Ситдикова Елена Анатольевна', 'tjde7c@yahoo.com', 'YOyhfR');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (10, 3, 'Ворсин Петр Евгеньевич', 'wpmrc3do@tutanota.com', 'RSbvHv');
INSERT INTO categories(id, name) VALUES (1, 'Прихожая');
INSERT INTO categories(id, name) VALUES (2, 'Диван');
INSERT INTO categories(id, name) VALUES (3, 'Обувница');
INSERT INTO categories(id, name) VALUES (4, 'Пуф');
INSERT INTO categories(id, name) VALUES (5, 'Полка');
INSERT INTO categories(id, name) VALUES (6, 'Стул');
INSERT INTO suppliers(id, name) VALUES (1, 'Стройландия');
INSERT INTO suppliers(id, name) VALUES (2, 'Кромма');
INSERT INTO suppliers(id, name) VALUES (3, 'ЗолотоеРуно');
INSERT INTO suppliers(id, name) VALUES (4, 'KRYLOVMANUFACTURA');
INSERT INTO manufacturers(id, name) VALUES (1, 'SVМЕБЕЛЬ');
INSERT INTO manufacturers(id, name) VALUES (2, 'Мебелони');
INSERT INTO manufacturers(id, name) VALUES (3, 'Инвуд');
INSERT INTO manufacturers(id, name) VALUES (4, 'RIDBERG');
INSERT INTO units(id, name) VALUES (1, 'шт.');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (1, 'А112Т4', 'Прихожая Фаворит 1 1420х2056х352ммм Дуб Делано/Цемент Светлый SV-М 1 шт', 1, 9577.0, 1, 1, 1, 10.0, 0, 'Удивительно функциональная и практичная прихожая Фаворит 1, обладая характерными чертами Скандинавского стиля, выглядит эффектно и способна задать тон интерьеру дома, встречая вас и ваших гостей.', 'assets/products/1.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (2, 'G843H5', 'Прихожая в коридор Твист с зеркалом мебель со шкафами, 120х37х202 см', 1, 8803.0, 1, 2, 1, 25.0, 9, 'Этот стеллаж со шкафом в прихожую комнату станет отличным элементом для вашего интерьера. Мебель для дома обеспечивает удобное хранение перчаток, шапок, зонтов, сумок и других аксессуаров.', 'assets/products/2.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (3, 'D325D4', 'Угловой диван Кромма Инвуд Лайт, серый двухместный диван, Velutto 32', 1, 29125.0, 2, 3, 2, 5.0, 12, 'Угловой диван Инвуд Лайт 2 - стильный и комфортный диван подойдет для комнаты любого размера.', 'assets/products/3.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (4, 'S432T5', 'Обувница RIDBERG, с вешалкой, стальная, 170x60x26 см, 5 полок, вместимость до 15 пар', 1, 885.0, 2, 4, 3, 15.0, 15, 'Обувница Ridberg с 5 полками и вешалкой - идеальное решение для организации хранения обуви в прихожей или гардеробной.', 'assets/products/4.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (5, 'F325D4', 'Диван, Прямой диван, Диван-кровать Сити темно-коричневый. Квест-33', 1, 14322.0, 3, 3, 2, 18.0, 3, 'Прямой диван-кровать Сити - это современное и функциональное решение для вашего дома.', 'assets/products/5.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (6, 'G432G6', 'Пуф трансформер кровать раскладушка светло-коричневый велюр', 1, 6149.0, 4, 3, 4, 22.0, 3, 'Пуф трансформер 5в1 представляет собой уникальное сочетание функций, выступая в качестве пуфика, столика, кресла, шезлонга и дополнительного спального места.', 'assets/products/6.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (7, 'H542F5', 'Диван, Прямой диван, диван кровать, Рио симпл механизм Пантограф. Симпл-16', 1, 20708.0, 3, 3, 2, 4.0, 5, 'Диван Рио симпл от "Золотое Руно" - это сочетание комфорта, функциональности и стильного дизайна.', 'assets/products/7.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (8, 'C346F5', 'Полка настенная ромб Лофт, черная, 40 см', 1, 2843.0, 4, 4, 5, 5.0, 4, 'Полочки для цветов в стиле лофт. Подойдут как для цветов, так и в качестве декоративного элемента. Полки подойдут для дома, офиса, кафе, ресторана. ​', 'assets/products/8.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (9, 'F256G6', 'Стулья для кухни', 1, 4760.0, 4, 4, 6, 6.0, 2, 'Набор из четырех стульев в лофт-дизайне станет любимой мебелью для отдыха и подойдет для взрослых и детей.', 'assets/products/9.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (10, 'J532V5', 'Магнитная полка, для холодильника, металл, 3шт, универсальная, чёрная', 1, 1387.0, 4, 4, 5, 8.0, 6, 'Магнитная полка для холодильника - это удобный и практичный аксессуар, который поможет организовать пространство в вашем доме.', 'assets/products/10.jpg');
INSERT INTO pickup_points(id, address) VALUES (1, '420151, г. Лесной, ул. Вишневая, 32');
INSERT INTO pickup_points(id, address) VALUES (2, '125061, г. Лесной, ул. Подгорная, 8');
INSERT INTO pickup_points(id, address) VALUES (3, '630370, г. Лесной, ул. Шоссейная, 24');
INSERT INTO pickup_points(id, address) VALUES (4, '400562, г. Лесной, ул. Зеленая, 32');
INSERT INTO pickup_points(id, address) VALUES (5, '614510, г. Лесной, ул. Маяковского, 47');
INSERT INTO pickup_points(id, address) VALUES (6, '410542, г. Лесной, ул. Светлая, 46');
INSERT INTO pickup_points(id, address) VALUES (7, '620839, г. Лесной, ул. Цветочная, 8');
INSERT INTO pickup_points(id, address) VALUES (8, '443890, г. Лесной, ул. Коммунистическая, 1');
INSERT INTO pickup_points(id, address) VALUES (9, '603379, г. Лесной, ул. Спортивная, 46');
INSERT INTO pickup_points(id, address) VALUES (10, '603721, г. Лесной, ул. Гоголя, 41');
INSERT INTO pickup_points(id, address) VALUES (11, '410172, г. Лесной, ул. Северная, 13');
INSERT INTO pickup_points(id, address) VALUES (12, '614611, г. Лесной, ул. Молодежная, 50');
INSERT INTO pickup_points(id, address) VALUES (13, '454311, г.Лесной, ул. Новая, 19');
INSERT INTO pickup_points(id, address) VALUES (14, '660007, г.Лесной, ул. Октябрьская, 19');
INSERT INTO pickup_points(id, address) VALUES (15, '603036, г. Лесной, ул. Садовая, 4');
INSERT INTO pickup_points(id, address) VALUES (16, '394060, г.Лесной, ул. Фрунзе, 43');
INSERT INTO pickup_points(id, address) VALUES (17, '410661, г. Лесной, ул. Школьная, 50');
INSERT INTO pickup_points(id, address) VALUES (18, '625590, г. Лесной, ул. Коммунистическая, 20');
INSERT INTO pickup_points(id, address) VALUES (19, '625683, г. Лесной, ул. 8 Марта');
INSERT INTO pickup_points(id, address) VALUES (20, '450983, г.Лесной, ул. Комсомольская, 26');
INSERT INTO pickup_points(id, address) VALUES (21, '394782, г. Лесной, ул. Чехова, 3');
INSERT INTO pickup_points(id, address) VALUES (22, '603002, г. Лесной, ул. Дзержинского, 28');
INSERT INTO pickup_points(id, address) VALUES (23, '450558, г. Лесной, ул. Набережная, 30');
INSERT INTO pickup_points(id, address) VALUES (24, '344288, г. Лесной, ул. Чехова, 1');
INSERT INTO pickup_points(id, address) VALUES (25, '614164, г.Лесной,  ул. Степная, 30');
INSERT INTO pickup_points(id, address) VALUES (26, '394242, г. Лесной, ул. Коммунистическая, 43');
INSERT INTO pickup_points(id, address) VALUES (27, '660540, г. Лесной, ул. Солнечная, 25');
INSERT INTO pickup_points(id, address) VALUES (28, '125837, г. Лесной, ул. Шоссейная, 40');
INSERT INTO pickup_points(id, address) VALUES (29, '125703, г. Лесной, ул. Партизанская, 49');
INSERT INTO pickup_points(id, address) VALUES (30, '625283, г. Лесной, ул. Победы, 46');
INSERT INTO pickup_points(id, address) VALUES (31, '614753, г. Лесной, ул. Полевая, 35');
INSERT INTO pickup_points(id, address) VALUES (32, '426030, г. Лесной, ул. Маяковского, 44');
INSERT INTO pickup_points(id, address) VALUES (33, '450375, г. Лесной ул. Клубная, 44');
INSERT INTO pickup_points(id, address) VALUES (34, '625560, г. Лесной, ул. Некрасова, 12');
INSERT INTO pickup_points(id, address) VALUES (35, '630201, г. Лесной, ул. Комсомольская, 17');
INSERT INTO pickup_points(id, address) VALUES (36, '190949, г. Лесной, ул. Мичурина, 26');
INSERT INTO order_statuses(id, name) VALUES (1, 'Новый');
INSERT INTO order_statuses(id, name) VALUES (2, 'Завершен');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (1, 1, 7, 'Степанов Михаил Артёмович', '901', 1, '2024-02-27', '2024-04-20');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (2, 11, 8, 'Михайлюк Анна Вячеславовна', '902', 1, '2024-09-28', '2024-04-21');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (3, 2, 9, 'Ситдикова Елена Анатольевна', '903', 1, '2024-03-21', '2024-04-22');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (4, 11, 10, 'Ворсин Петр Евгеньевич', '904', 2, '2024-02-20', '2024-04-23');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (5, 2, 7, 'Степанов Михаил Артёмович', '905', 2, '2024-03-17', '2024-04-24');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (6, 15, 8, 'Михайлюк Анна Вячеславовна', '906', 2, '2024-03-01', '2024-04-25');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (7, 3, 9, 'Ситдикова Елена Анатольевна', '907', 2, NULL, '2024-04-26');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (8, 19, 10, 'Ворсин Петр Евгеньевич', '908', 1, '2024-03-31', '2024-04-27');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (9, 5, 9, 'Ситдикова Елена Анатольевна', '909', 1, '2024-04-02', '2024-04-28');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (10, 19, 10, 'Ворсин Петр Евгеньевич', '910', 1, '2024-04-03', '2024-04-29');
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (1, 1, 1, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (2, 1, 2, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (3, 2, 2, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (4, 2, 1, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (5, 3, 3, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (6, 3, 4, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (7, 4, 5, 5);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (8, 4, 3, 4);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (9, 5, 6, 20);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (10, 5, 7, 20);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (11, 6, 1, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (12, 6, 2, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (13, 7, 2, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (14, 7, 1, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (15, 8, 3, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (16, 8, 4, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (17, 9, 5, 5);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (18, 9, 3, 4);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (19, 10, 6, 20);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (20, 10, 7, 20);
''',
    # ============================================================
    # END FILE: database/seed_sqlite.sql
    # ============================================================

    # ============================================================
    # BEGIN FILE: database/schema_mysql.sql
    # ============================================================
    'database/schema_mysql.sql': r'''CREATE TABLE IF NOT EXISTS roles (
    id INT PRIMARY KEY AUTO_INCREMENT,
    name VARCHAR(100) NOT NULL UNIQUE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS users (
    id INT PRIMARY KEY AUTO_INCREMENT,
    role_id INT NOT NULL,
    full_name VARCHAR(255) NOT NULL,
    login VARCHAR(255) NOT NULL UNIQUE,
    password VARCHAR(100) NOT NULL,
    CONSTRAINT fk_users_roles FOREIGN KEY (role_id) REFERENCES roles(id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS categories (
    id INT PRIMARY KEY AUTO_INCREMENT,
    name VARCHAR(255) NOT NULL UNIQUE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS suppliers (
    id INT PRIMARY KEY AUTO_INCREMENT,
    name VARCHAR(255) NOT NULL UNIQUE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS manufacturers (
    id INT PRIMARY KEY AUTO_INCREMENT,
    name VARCHAR(255) NOT NULL UNIQUE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS units (
    id INT PRIMARY KEY AUTO_INCREMENT,
    name VARCHAR(100) NOT NULL UNIQUE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS products (
    id INT PRIMARY KEY AUTO_INCREMENT,
    article VARCHAR(50) NOT NULL UNIQUE,
    name VARCHAR(500) NOT NULL,
    unit_id INT NOT NULL,
    price DECIMAL(10, 2) NOT NULL,
    supplier_id INT NOT NULL,
    manufacturer_id INT NOT NULL,
    category_id INT NOT NULL,
    discount DECIMAL(5, 2) NOT NULL DEFAULT 0,
    stock_quantity INT NOT NULL DEFAULT 0,
    description TEXT NOT NULL,
    photo_path VARCHAR(500),
    CONSTRAINT chk_products_price CHECK (price >= 0),
    CONSTRAINT chk_products_discount CHECK (discount >= 0 AND discount <= 100),
    CONSTRAINT chk_products_stock CHECK (stock_quantity >= 0),
    CONSTRAINT fk_products_units FOREIGN KEY (unit_id) REFERENCES units(id),
    CONSTRAINT fk_products_suppliers FOREIGN KEY (supplier_id) REFERENCES suppliers(id),
    CONSTRAINT fk_products_manufacturers FOREIGN KEY (manufacturer_id) REFERENCES manufacturers(id),
    CONSTRAINT fk_products_categories FOREIGN KEY (category_id) REFERENCES categories(id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS pickup_points (
    id INT PRIMARY KEY AUTO_INCREMENT,
    address VARCHAR(500) NOT NULL UNIQUE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS order_statuses (
    id INT PRIMARY KEY AUTO_INCREMENT,
    name VARCHAR(100) NOT NULL UNIQUE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS orders (
    id INT PRIMARY KEY AUTO_INCREMENT,
    pickup_point_id INT NOT NULL,
    client_user_id INT,
    client_name VARCHAR(255),
    receive_code VARCHAR(50),
    status_id INT NOT NULL,
    order_date DATE,
    delivery_date DATE,
    CONSTRAINT fk_orders_pickup_points FOREIGN KEY (pickup_point_id) REFERENCES pickup_points(id),
    CONSTRAINT fk_orders_users FOREIGN KEY (client_user_id) REFERENCES users(id),
    CONSTRAINT fk_orders_statuses FOREIGN KEY (status_id) REFERENCES order_statuses(id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;

CREATE TABLE IF NOT EXISTS order_items (
    id INT PRIMARY KEY AUTO_INCREMENT,
    order_id INT NOT NULL,
    product_id INT NOT NULL,
    quantity INT NOT NULL,
    CONSTRAINT chk_order_items_quantity CHECK (quantity > 0),
    CONSTRAINT fk_order_items_orders FOREIGN KEY (order_id) REFERENCES orders(id) ON DELETE CASCADE,
    CONSTRAINT fk_order_items_products FOREIGN KEY (product_id) REFERENCES products(id),
    CONSTRAINT uq_order_items_order_product UNIQUE (order_id, product_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;
''',
    # ============================================================
    # END FILE: database/schema_mysql.sql
    # ============================================================

    # ============================================================
    # BEGIN FILE: database/seed_mysql.sql
    # ============================================================
    'database/seed_mysql.sql': r'''-- Начальные данные для демонстрационного экзамена, вариант 5.
INSERT INTO roles(id, name) VALUES (1, 'Администратор');
INSERT INTO roles(id, name) VALUES (2, 'Менеджер');
INSERT INTO roles(id, name) VALUES (3, 'Авторизированный клиент');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (1, 1, 'Никифорова Анна Семеновна', '94d5ous@gmail.com', 'uzWC67');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (2, 1, 'Стелина Евгения Петровна', 'uth4iz@mail.com', '2L6KZG');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (3, 1, 'Никифорова Весения Николаевна', '5d4zbu@tutanota.com', 'rwVDh9');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (4, 2, 'Сазонов Руслан Германович', 'ptec8ym@yahoo.com', 'LdNyos');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (5, 2, 'Одинцов Серафим Артёмович', '1qz4kw@mail.com', 'gynQMT');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (6, 2, 'Старикова Елена Павловна', '4np6se@mail.com', 'AtnDjr');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (7, 3, 'Степанов Михаил Артёмович', 'yzls62@outlook.com', 'JlFRCZ');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (8, 3, 'Михайлюк Анна Вячеславовна', '1diph5e@tutanota.com', '8ntwUp');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (9, 3, 'Ситдикова Елена Анатольевна', 'tjde7c@yahoo.com', 'YOyhfR');
INSERT INTO users(id, role_id, full_name, login, password) VALUES (10, 3, 'Ворсин Петр Евгеньевич', 'wpmrc3do@tutanota.com', 'RSbvHv');
INSERT INTO categories(id, name) VALUES (1, 'Прихожая');
INSERT INTO categories(id, name) VALUES (2, 'Диван');
INSERT INTO categories(id, name) VALUES (3, 'Обувница');
INSERT INTO categories(id, name) VALUES (4, 'Пуф');
INSERT INTO categories(id, name) VALUES (5, 'Полка');
INSERT INTO categories(id, name) VALUES (6, 'Стул');
INSERT INTO suppliers(id, name) VALUES (1, 'Стройландия');
INSERT INTO suppliers(id, name) VALUES (2, 'Кромма');
INSERT INTO suppliers(id, name) VALUES (3, 'ЗолотоеРуно');
INSERT INTO suppliers(id, name) VALUES (4, 'KRYLOVMANUFACTURA');
INSERT INTO manufacturers(id, name) VALUES (1, 'SVМЕБЕЛЬ');
INSERT INTO manufacturers(id, name) VALUES (2, 'Мебелони');
INSERT INTO manufacturers(id, name) VALUES (3, 'Инвуд');
INSERT INTO manufacturers(id, name) VALUES (4, 'RIDBERG');
INSERT INTO units(id, name) VALUES (1, 'шт.');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (1, 'А112Т4', 'Прихожая Фаворит 1 1420х2056х352ммм Дуб Делано/Цемент Светлый SV-М 1 шт', 1, 9577.0, 1, 1, 1, 10.0, 0, 'Удивительно функциональная и практичная прихожая Фаворит 1, обладая характерными чертами Скандинавского стиля, выглядит эффектно и способна задать тон интерьеру дома, встречая вас и ваших гостей.', 'assets/products/1.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (2, 'G843H5', 'Прихожая в коридор Твист с зеркалом мебель со шкафами, 120х37х202 см', 1, 8803.0, 1, 2, 1, 25.0, 9, 'Этот стеллаж со шкафом в прихожую комнату станет отличным элементом для вашего интерьера. Мебель для дома обеспечивает удобное хранение перчаток, шапок, зонтов, сумок и других аксессуаров.', 'assets/products/2.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (3, 'D325D4', 'Угловой диван Кромма Инвуд Лайт, серый двухместный диван, Velutto 32', 1, 29125.0, 2, 3, 2, 5.0, 12, 'Угловой диван Инвуд Лайт 2 - стильный и комфортный диван подойдет для комнаты любого размера.', 'assets/products/3.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (4, 'S432T5', 'Обувница RIDBERG, с вешалкой, стальная, 170x60x26 см, 5 полок, вместимость до 15 пар', 1, 885.0, 2, 4, 3, 15.0, 15, 'Обувница Ridberg с 5 полками и вешалкой - идеальное решение для организации хранения обуви в прихожей или гардеробной.', 'assets/products/4.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (5, 'F325D4', 'Диван, Прямой диван, Диван-кровать Сити темно-коричневый. Квест-33', 1, 14322.0, 3, 3, 2, 18.0, 3, 'Прямой диван-кровать Сити - это современное и функциональное решение для вашего дома.', 'assets/products/5.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (6, 'G432G6', 'Пуф трансформер кровать раскладушка светло-коричневый велюр', 1, 6149.0, 4, 3, 4, 22.0, 3, 'Пуф трансформер 5в1 представляет собой уникальное сочетание функций, выступая в качестве пуфика, столика, кресла, шезлонга и дополнительного спального места.', 'assets/products/6.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (7, 'H542F5', 'Диван, Прямой диван, диван кровать, Рио симпл механизм Пантограф. Симпл-16', 1, 20708.0, 3, 3, 2, 4.0, 5, 'Диван Рио симпл от "Золотое Руно" - это сочетание комфорта, функциональности и стильного дизайна.', 'assets/products/7.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (8, 'C346F5', 'Полка настенная ромб Лофт, черная, 40 см', 1, 2843.0, 4, 4, 5, 5.0, 4, 'Полочки для цветов в стиле лофт. Подойдут как для цветов, так и в качестве декоративного элемента. Полки подойдут для дома, офиса, кафе, ресторана. ​', 'assets/products/8.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (9, 'F256G6', 'Стулья для кухни', 1, 4760.0, 4, 4, 6, 6.0, 2, 'Набор из четырех стульев в лофт-дизайне станет любимой мебелью для отдыха и подойдет для взрослых и детей.', 'assets/products/9.jpg');
INSERT INTO products(id, article, name, unit_id, price, supplier_id, manufacturer_id, category_id, discount, stock_quantity, description, photo_path) VALUES (10, 'J532V5', 'Магнитная полка, для холодильника, металл, 3шт, универсальная, чёрная', 1, 1387.0, 4, 4, 5, 8.0, 6, 'Магнитная полка для холодильника - это удобный и практичный аксессуар, который поможет организовать пространство в вашем доме.', 'assets/products/10.jpg');
INSERT INTO pickup_points(id, address) VALUES (1, '420151, г. Лесной, ул. Вишневая, 32');
INSERT INTO pickup_points(id, address) VALUES (2, '125061, г. Лесной, ул. Подгорная, 8');
INSERT INTO pickup_points(id, address) VALUES (3, '630370, г. Лесной, ул. Шоссейная, 24');
INSERT INTO pickup_points(id, address) VALUES (4, '400562, г. Лесной, ул. Зеленая, 32');
INSERT INTO pickup_points(id, address) VALUES (5, '614510, г. Лесной, ул. Маяковского, 47');
INSERT INTO pickup_points(id, address) VALUES (6, '410542, г. Лесной, ул. Светлая, 46');
INSERT INTO pickup_points(id, address) VALUES (7, '620839, г. Лесной, ул. Цветочная, 8');
INSERT INTO pickup_points(id, address) VALUES (8, '443890, г. Лесной, ул. Коммунистическая, 1');
INSERT INTO pickup_points(id, address) VALUES (9, '603379, г. Лесной, ул. Спортивная, 46');
INSERT INTO pickup_points(id, address) VALUES (10, '603721, г. Лесной, ул. Гоголя, 41');
INSERT INTO pickup_points(id, address) VALUES (11, '410172, г. Лесной, ул. Северная, 13');
INSERT INTO pickup_points(id, address) VALUES (12, '614611, г. Лесной, ул. Молодежная, 50');
INSERT INTO pickup_points(id, address) VALUES (13, '454311, г.Лесной, ул. Новая, 19');
INSERT INTO pickup_points(id, address) VALUES (14, '660007, г.Лесной, ул. Октябрьская, 19');
INSERT INTO pickup_points(id, address) VALUES (15, '603036, г. Лесной, ул. Садовая, 4');
INSERT INTO pickup_points(id, address) VALUES (16, '394060, г.Лесной, ул. Фрунзе, 43');
INSERT INTO pickup_points(id, address) VALUES (17, '410661, г. Лесной, ул. Школьная, 50');
INSERT INTO pickup_points(id, address) VALUES (18, '625590, г. Лесной, ул. Коммунистическая, 20');
INSERT INTO pickup_points(id, address) VALUES (19, '625683, г. Лесной, ул. 8 Марта');
INSERT INTO pickup_points(id, address) VALUES (20, '450983, г.Лесной, ул. Комсомольская, 26');
INSERT INTO pickup_points(id, address) VALUES (21, '394782, г. Лесной, ул. Чехова, 3');
INSERT INTO pickup_points(id, address) VALUES (22, '603002, г. Лесной, ул. Дзержинского, 28');
INSERT INTO pickup_points(id, address) VALUES (23, '450558, г. Лесной, ул. Набережная, 30');
INSERT INTO pickup_points(id, address) VALUES (24, '344288, г. Лесной, ул. Чехова, 1');
INSERT INTO pickup_points(id, address) VALUES (25, '614164, г.Лесной,  ул. Степная, 30');
INSERT INTO pickup_points(id, address) VALUES (26, '394242, г. Лесной, ул. Коммунистическая, 43');
INSERT INTO pickup_points(id, address) VALUES (27, '660540, г. Лесной, ул. Солнечная, 25');
INSERT INTO pickup_points(id, address) VALUES (28, '125837, г. Лесной, ул. Шоссейная, 40');
INSERT INTO pickup_points(id, address) VALUES (29, '125703, г. Лесной, ул. Партизанская, 49');
INSERT INTO pickup_points(id, address) VALUES (30, '625283, г. Лесной, ул. Победы, 46');
INSERT INTO pickup_points(id, address) VALUES (31, '614753, г. Лесной, ул. Полевая, 35');
INSERT INTO pickup_points(id, address) VALUES (32, '426030, г. Лесной, ул. Маяковского, 44');
INSERT INTO pickup_points(id, address) VALUES (33, '450375, г. Лесной ул. Клубная, 44');
INSERT INTO pickup_points(id, address) VALUES (34, '625560, г. Лесной, ул. Некрасова, 12');
INSERT INTO pickup_points(id, address) VALUES (35, '630201, г. Лесной, ул. Комсомольская, 17');
INSERT INTO pickup_points(id, address) VALUES (36, '190949, г. Лесной, ул. Мичурина, 26');
INSERT INTO order_statuses(id, name) VALUES (1, 'Новый');
INSERT INTO order_statuses(id, name) VALUES (2, 'Завершен');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (1, 1, 7, 'Степанов Михаил Артёмович', '901', 1, '2024-02-27', '2024-04-20');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (2, 11, 8, 'Михайлюк Анна Вячеславовна', '902', 1, '2024-09-28', '2024-04-21');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (3, 2, 9, 'Ситдикова Елена Анатольевна', '903', 1, '2024-03-21', '2024-04-22');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (4, 11, 10, 'Ворсин Петр Евгеньевич', '904', 2, '2024-02-20', '2024-04-23');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (5, 2, 7, 'Степанов Михаил Артёмович', '905', 2, '2024-03-17', '2024-04-24');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (6, 15, 8, 'Михайлюк Анна Вячеславовна', '906', 2, '2024-03-01', '2024-04-25');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (7, 3, 9, 'Ситдикова Елена Анатольевна', '907', 2, NULL, '2024-04-26');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (8, 19, 10, 'Ворсин Петр Евгеньевич', '908', 1, '2024-03-31', '2024-04-27');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (9, 5, 9, 'Ситдикова Елена Анатольевна', '909', 1, '2024-04-02', '2024-04-28');
INSERT INTO orders(id, pickup_point_id, client_user_id, client_name, receive_code, status_id, order_date, delivery_date) VALUES (10, 19, 10, 'Ворсин Петр Евгеньевич', '910', 1, '2024-04-03', '2024-04-29');
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (1, 1, 1, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (2, 1, 2, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (3, 2, 2, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (4, 2, 1, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (5, 3, 3, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (6, 3, 4, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (7, 4, 5, 5);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (8, 4, 3, 4);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (9, 5, 6, 20);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (10, 5, 7, 20);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (11, 6, 1, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (12, 6, 2, 2);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (13, 7, 2, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (14, 7, 1, 1);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (15, 8, 3, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (16, 8, 4, 10);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (17, 9, 5, 5);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (18, 9, 3, 4);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (19, 10, 6, 20);
INSERT INTO order_items(id, order_id, product_id, quantity) VALUES (20, 10, 7, 20);
''',
    # ============================================================
    # END FILE: database/seed_mysql.sql
    # ============================================================

    # ============================================================
    # BEGIN FILE: database/import_report.txt
    # ============================================================
    'database/import_report.txt': r'''Отчет о мягкой очистке import-данных

- Заказ 7: некорректная дата 30.02.2024 заменена пустым значением.
''',
    # ============================================================
    # END FILE: database/import_report.txt
    # ============================================================

    # ============================================================
    # BEGIN FILE: tools/generate_artifacts.py
    # ============================================================
    'tools/generate_artifacts.py': r'''import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A3, A4, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


PROJECT_ROOT = Path(__file__).resolve().parents[1]
APP_DIR = PROJECT_ROOT / "app"
DOCS_DIR = PROJECT_ROOT / "docs"
SCREENSHOTS_DIR = DOCS_DIR / "screenshots"
sys.path.insert(0, str(APP_DIR))

from config import (
    COLOR_ACCENT,
    COLOR_DISCOUNT,
    COLOR_OUT_OF_STOCK,
    COLOR_PANEL,
    COLOR_SECONDARY,
    COLOR_WHITE,
    LOGO_PATH,
)
from database import create_sql_dump, initialize_database
from repositories import OrderRepository, ProductRepository, resolve_asset_path


FONT_PATH = Path("C:/Windows/Fonts/calibri.ttf")
BOLD_FONT_PATH = Path("C:/Windows/Fonts/calibrib.ttf")
PDF_FONT_NAME = "Calibri"


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = BOLD_FONT_PATH if bold and BOLD_FONT_PATH.exists() else FONT_PATH
    if path.exists():
        return ImageFont.truetype(str(path), size)

    return ImageFont.load_default()


def register_pdf_font() -> None:
    global PDF_FONT_NAME

    if FONT_PATH.exists():
        pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(FONT_PATH)))
        return

    PDF_FONT_NAME = "Helvetica"


def draw_wrapped_text(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], font, fill, max_width: int, line_spacing: int = 4) -> int:
    words = str(text).split()
    lines = []
    current = ""

    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word

    if current:
        lines.append(current)

    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_spacing

    return y


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], fill: str, outline: str = "#C7C7C7") -> None:
    draw.rounded_rectangle(xy, radius=6, fill=fill, outline=outline, width=1)


def open_or_make_placeholder(path: Path, size: tuple[int, int]) -> Image.Image:
    if path.exists():
        image = Image.open(path).convert("RGB")
        image.thumbnail(size)
        return image

    image = Image.new("RGB", size, COLOR_PANEL)
    draw = ImageDraw.Draw(image)
    draw.rectangle((3, 3, size[0] - 4, size[1] - 4), outline=COLOR_SECONDARY, width=2)
    draw.text((12, size[1] // 2 - 8), "Нет фото", font=get_font(14, True), fill=COLOR_ACCENT)
    return image


def create_login_screenshot(output_path: Path) -> None:
    image = Image.new("RGB", (1200, 760), COLOR_WHITE)
    draw = ImageDraw.Draw(image)

    title_font = get_font(34, True)
    label_font = get_font(18, True)
    text_font = get_font(18)

    draw.rectangle((0, 0, 1200, 760), fill=COLOR_WHITE)
    draw_box(draw, (390, 95, 810, 665), "#FFFFFF", "#80FFFF")

    if LOGO_PATH.exists():
        logo = Image.open(LOGO_PATH).convert("RGBA")
        logo.thumbnail((280, 165))
        image.paste(logo, (600 - logo.width // 2, 125), logo)

    draw.text((492, 310), "Вход в систему", font=title_font, fill=COLOR_ACCENT)
    draw.text((440, 385), "Логин", font=label_font, fill="#222222")
    draw_box(draw, (440, 415, 760, 462), "#FFFFFF", "#808080")
    draw.text((455, 425), "94d5ous@gmail.com", font=text_font, fill="#333333")
    draw.text((440, 480), "Пароль", font=label_font, fill="#222222")
    draw_box(draw, (440, 510, 760, 557), "#FFFFFF", "#808080")
    draw.text((455, 520), "••••••", font=text_font, fill="#333333")
    draw_box(draw, (440, 585, 760, 635), COLOR_ACCENT, COLOR_ACCENT)
    draw.text((565, 598), "Войти", font=label_font, fill=COLOR_WHITE)

    image.save(output_path)


def draw_product_card(draw: ImageDraw.ImageDraw, image: Image.Image, product: dict, y: int, mode: str = "default") -> int:
    background = "#FFFFFF"
    foreground = "#111111"
    if product["stock_quantity"] == 0:
        background = COLOR_OUT_OF_STOCK
    elif product["discount"] > 15:
        background = COLOR_DISCOUNT
        foreground = "#FFFFFF"

    draw_box(draw, (45, y, 1155, y + 142), background, "#BDBDBD")
    photo_path = resolve_asset_path(product.get("photo_path"))
    photo = open_or_make_placeholder(photo_path, (130, 95))
    image.paste(photo, (65, y + 24))

    title_font = get_font(18, True)
    text_font = get_font(15)
    small_font = get_font(14)

    draw_wrapped_text(draw, f"{product['article']} - {product['name']}", (220, y + 16), title_font, foreground, 700)
    draw.text((220, y + 64), f"Категория: {product['category_name']} | Производитель: {product['manufacturer_name']}", font=text_font, fill=foreground)
    draw.text((220, y + 91), f"Поставщик: {product['supplier_name']} | Остаток: {product['stock_quantity']} | Скидка: {product['discount']:.0f}%", font=text_font, fill=foreground)

    price_x = 890 if mode == "admin" else 930
    if product["discount"] > 0:
        draw.text((price_x, y + 46), f"{product['price']:.2f} руб.", font=small_font, fill="red")
        draw.line((price_x, y + 57, price_x + 108, y + 57), fill="red", width=2)
        draw.text((price_x, y + 72), f"{product['final_price']:.2f} руб.", font=title_font, fill="#000000")
    else:
        draw.text((930, y + 58), f"{product['price']:.2f} руб.", font=title_font, fill=foreground)

    if mode == "admin":
        draw_box(draw, (1035, y + 18, 1135, y + 54), "#F7F7F7", "#888888")
        draw.text((1045, y + 26), "Изменить", font=small_font, fill="#111111")
        draw_box(draw, (1035, y + 70, 1135, y + 106), "#F7F7F7", "#888888")
        draw.text((1055, y + 78), "Удалить", font=small_font, fill="#111111")

    return y + 160


def create_products_screenshot(output_path: Path, role: str, mode: str) -> None:
    repository = ProductRepository()
    if mode == "manager":
        products = repository.list_products("диван", "15% и более", "Цена", "DESC")
    else:
        products = repository.list_products()

    image = Image.new("RGB", (1200, 760), COLOR_WHITE)
    draw = ImageDraw.Draw(image)
    title_font = get_font(26, True)
    text_font = get_font(16)
    small_font = get_font(14)

    draw.rectangle((0, 0, 1200, 64), fill=COLOR_SECONDARY)
    draw.text((25, 18), "Список товаров", font=title_font, fill="#111111")
    draw.text((760, 22), role[:34], font=text_font, fill="#111111")
    draw_box(draw, (1085, 16, 1160, 48), "#FFFFFF", "#FFFFFF")
    draw.text((1102, 23), "Выйти", font=small_font, fill="#111111")

    if mode in {"manager", "admin"}:
        draw_box(draw, (35, 88, 1165, 155), COLOR_PANEL, "#80FFFF")
        draw.text((55, 102), "Поиск: диван", font=text_font, fill="#111111")
        draw.text((275, 102), "Скидка: 15% и более", font=text_font, fill="#111111")
        draw.text((520, 102), "Сортировка: Цена DESC", font=text_font, fill="#111111")
        draw_box(draw, (965, 103, 1058, 138), COLOR_ACCENT, COLOR_ACCENT)
        draw.text((981, 111), "Заказы", font=small_font, fill=COLOR_WHITE)
        if mode == "admin":
            draw_box(draw, (815, 103, 945, 138), COLOR_ACCENT, COLOR_ACCENT)
            draw.text((829, 111), "Добавить товар", font=small_font, fill=COLOR_WHITE)
        y = 175
    else:
        y = 90

    for product in products[:4]:
        y = draw_product_card(draw, image, product, y, mode)

    image.save(output_path)


def create_product_form_screenshot(output_path: Path) -> None:
    product = ProductRepository().list_products()[0]
    image = Image.new("RGB", (1200, 760), COLOR_WHITE)
    draw = ImageDraw.Draw(image)
    title_font = get_font(28, True)
    label_font = get_font(16, True)
    text_font = get_font(16)

    draw.text((45, 35), "Редактирование товара", font=title_font, fill=COLOR_ACCENT)
    fields = [
        ("ID", str(product["id"])),
        ("Артикул", product["article"]),
        ("Наименование товара", product["name"][:72]),
        ("Категория товара", product["category_name"]),
        ("Производитель", product["manufacturer_name"]),
        ("Поставщик", product["supplier_name"]),
        ("Цена", f"{product['price']:.2f}"),
        ("Единица измерения", product["unit_name"]),
        ("Количество на складе", str(product["stock_quantity"])),
        ("Действующая скидка", f"{product['discount']:.2f}"),
    ]

    y = 95
    for label, value in fields:
        draw.text((45, y + 11), label, font=label_font, fill="#111111")
        draw_box(draw, (280, y, 880, y + 42), "#FFFFFF", "#999999")
        draw.text((295, y + 10), value, font=text_font, fill="#333333")
        y += 54

    draw.text((45, y + 10), "Описание товара", font=label_font, fill="#111111")
    draw_box(draw, (280, y, 880, y + 105), "#FFFFFF", "#999999")
    draw_wrapped_text(draw, product["description"], (295, y + 10), text_font, "#333333", 560)

    photo_path = resolve_asset_path(product.get("photo_path"))
    photo = open_or_make_placeholder(photo_path, (220, 150))
    image.paste(photo, (930, 170))
    draw_box(draw, (930, 350, 1110, 390), COLOR_ACCENT, COLOR_ACCENT)
    draw.text((960, 360), "Выбрать фото", font=text_font, fill=COLOR_WHITE)
    draw_box(draw, (820, 685, 960, 728), COLOR_ACCENT, COLOR_ACCENT)
    draw.text((848, 696), "Сохранить", font=text_font, fill=COLOR_WHITE)

    image.save(output_path)


def create_orders_screenshot(output_path: Path) -> None:
    orders = OrderRepository().list_orders()
    image = Image.new("RGB", (1200, 760), COLOR_WHITE)
    draw = ImageDraw.Draw(image)
    title_font = get_font(26, True)
    header_font = get_font(15, True)
    text_font = get_font(13)

    draw.rectangle((0, 0, 1200, 64), fill=COLOR_SECONDARY)
    draw.text((25, 18), "Заказы", font=title_font, fill="#111111")
    draw.text((820, 22), "Никифорова Весения Николаевна", font=get_font(16), fill="#111111")
    draw_box(draw, (35, 88, 1165, 145), COLOR_PANEL, "#80FFFF")
    draw.text((55, 106), "Назад к товарам", font=get_font(16), fill="#111111")
    draw_box(draw, (230, 101, 360, 136), COLOR_ACCENT, COLOR_ACCENT)
    draw.text((244, 109), "Добавить заказ", font=get_font(13), fill=COLOR_WHITE)

    columns = [
        ("Номер", 55, 75),
        ("Статус", 130, 115),
        ("Пункт выдачи", 245, 280),
        ("Дата заказа", 525, 115),
        ("Дата выдачи", 640, 115),
        ("Клиент", 755, 220),
        ("Артикулы", 975, 170),
    ]
    y = 170
    x = 45
    for title, _start, width in columns:
        draw.rectangle((x, y, x + width, y + 34), fill=COLOR_ACCENT, outline="#FFFFFF")
        draw.text((x + 6, y + 9), title, font=header_font, fill=COLOR_WHITE)
        x += width

    y += 34
    for order in orders[:9]:
        x = 45
        values = [
            order["id"],
            order["status_name"],
            order["pickup_address"][:35],
            order["order_date"] or "",
            order["delivery_date"] or "",
            (order["client_name"] or "")[:28],
            (order["items_summary"] or "")[:24],
        ]
        for value, column in zip(values, columns):
            width = column[2]
            draw.rectangle((x, y, x + width, y + 44), fill="#FFFFFF", outline="#D0D0D0")
            draw.text((x + 6, y + 13), str(value), font=text_font, fill="#111111")
            x += width
        y += 44

    image.save(output_path)


def create_order_form_screenshot(output_path: Path) -> None:
    order = OrderRepository().get_order(1)
    image = Image.new("RGB", (1200, 760), COLOR_WHITE)
    draw = ImageDraw.Draw(image)
    title_font = get_font(28, True)
    label_font = get_font(16, True)
    text_font = get_font(16)

    draw.text((45, 35), "Редактирование заказа", font=title_font, fill=COLOR_ACCENT)
    fields = [
        ("Номер заказа", str(order["id"])),
        ("Статус заказа", "Новый"),
        ("Адрес пункта выдачи", "1 - 420151, г. Лесной, ул. Вишневая, 32"),
        ("ФИО клиента", order["client_name"]),
        ("Код для получения", order["receive_code"]),
        ("Дата заказа", order["order_date"] or ""),
        ("Дата выдачи", order["delivery_date"] or ""),
    ]

    y = 95
    for label, value in fields:
        draw.text((45, y + 11), label, font=label_font, fill="#111111")
        draw_box(draw, (290, y, 1050, y + 42), "#FFFFFF", "#999999")
        draw.text((305, y + 10), str(value), font=text_font, fill="#333333")
        y += 54

    draw.text((45, y + 20), "Позиции заказа", font=label_font, fill="#111111")
    y += 55
    draw.rectangle((45, y, 1050, y + 38), fill=COLOR_ACCENT)
    draw.text((65, y + 10), "Артикул", font=label_font, fill=COLOR_WHITE)
    draw.text((260, y + 10), "Товар", font=label_font, fill=COLOR_WHITE)
    draw.text((910, y + 10), "Количество", font=label_font, fill=COLOR_WHITE)
    y += 38

    for item in order["items"]:
        draw.rectangle((45, y, 1050, y + 48), fill="#FFFFFF", outline="#D0D0D0")
        draw.text((65, y + 14), item["article"], font=text_font, fill="#111111")
        draw.text((260, y + 14), item["name"][:70], font=text_font, fill="#111111")
        draw.text((930, y + 14), str(item["quantity"]), font=text_font, fill="#111111")
        y += 48

    draw_box(draw, (880, 685, 1020, 728), COLOR_ACCENT, COLOR_ACCENT)
    draw.text((908, 696), "Сохранить", font=text_font, fill=COLOR_WHITE)
    image.save(output_path)


def generate_screenshots() -> list[Path]:
    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)
    screenshots = [
        (SCREENSHOTS_DIR / "01_login.png", create_login_screenshot),
        (SCREENSHOTS_DIR / "02_products_guest.png", lambda path: create_products_screenshot(path, "Гость", "guest")),
        (SCREENSHOTS_DIR / "03_products_manager_filter.png", lambda path: create_products_screenshot(path, "Ситдикова Елена Анатольевна", "manager")),
        (SCREENSHOTS_DIR / "04_products_admin.png", lambda path: create_products_screenshot(path, "Никифорова Весения Николаевна", "admin")),
        (SCREENSHOTS_DIR / "05_product_form.png", create_product_form_screenshot),
        (SCREENSHOTS_DIR / "06_orders_admin.png", create_orders_screenshot),
        (SCREENSHOTS_DIR / "07_order_form.png", create_order_form_screenshot),
    ]

    for path, factory in screenshots:
        factory(path)

    return [path for path, _factory in screenshots]


def draw_pdf_table(c: canvas.Canvas, name: str, fields: list[str], x: int, y: int, width: int) -> tuple[int, int]:
    row_height = 18
    height = 26 + row_height * len(fields)
    c.setFillColor(colors.HexColor(COLOR_ACCENT))
    c.rect(x, y - height, width, height, fill=True, stroke=False)
    c.setFillColor(colors.white)
    c.setFont(PDF_FONT_NAME, 11)
    c.drawString(x + 8, y - 17, name)
    c.setFillColor(colors.white)
    c.rect(x, y - height, width, height - 26, fill=True, stroke=True)
    c.setFillColor(colors.black)
    c.setFont(PDF_FONT_NAME, 8)

    text_y = y - 40
    for field in fields:
        c.drawString(x + 8, text_y, field)
        text_y -= row_height

    return x + width // 2, y - height // 2


def generate_er_pdf(output_path: Path) -> None:
    register_pdf_font()
    c = canvas.Canvas(str(output_path), pagesize=landscape(A3))
    c.setTitle("ER-диаграмма МебельОрг")
    c.setFont(PDF_FONT_NAME, 20)
    c.drawString(40, 795, "ER-диаграмма базы данных ООО «МебельОрг»")

    tables = {
        "roles": ["PK id", "name"],
        "users": ["PK id", "FK role_id", "full_name", "login", "password"],
        "categories": ["PK id", "name"],
        "suppliers": ["PK id", "name"],
        "manufacturers": ["PK id", "name"],
        "units": ["PK id", "name"],
        "products": ["PK id", "article", "name", "FK unit_id", "price", "FK supplier_id", "FK manufacturer_id", "FK category_id", "discount", "stock_quantity", "description", "photo_path"],
        "pickup_points": ["PK id", "address"],
        "order_statuses": ["PK id", "name"],
        "orders": ["PK id", "FK pickup_point_id", "FK client_user_id", "client_name", "receive_code", "FK status_id", "order_date", "delivery_date"],
        "order_items": ["PK id", "FK order_id", "FK product_id", "quantity"],
    }
    positions = {
        "roles": (40, 740, 160),
        "users": (240, 740, 210),
        "categories": (500, 740, 170),
        "suppliers": (710, 740, 170),
        "manufacturers": (920, 740, 190),
        "units": (40, 440, 160),
        "products": (240, 520, 260),
        "pickup_points": (540, 440, 220),
        "order_statuses": (800, 440, 190),
        "orders": (540, 250, 260),
        "order_items": (840, 250, 220),
    }

    centers = {}
    for table_name, fields in tables.items():
        centers[table_name] = draw_pdf_table(c, table_name, fields, *positions[table_name])

    relations = [
        ("users", "roles", "role_id"),
        ("products", "categories", "category_id"),
        ("products", "suppliers", "supplier_id"),
        ("products", "manufacturers", "manufacturer_id"),
        ("products", "units", "unit_id"),
        ("orders", "pickup_points", "pickup_point_id"),
        ("orders", "users", "client_user_id"),
        ("orders", "order_statuses", "status_id"),
        ("order_items", "orders", "order_id"),
        ("order_items", "products", "product_id"),
    ]

    c.setStrokeColor(colors.HexColor(COLOR_SECONDARY))
    c.setFillColor(colors.HexColor(COLOR_SECONDARY))
    c.setFont(PDF_FONT_NAME, 7)
    for child, parent, label in relations:
        x1, y1 = centers[child]
        x2, y2 = centers[parent]
        c.line(x1, y1, x2, y2)
        c.drawString((x1 + x2) / 2, (y1 + y2) / 2 + 4, label)

    c.showPage()
    c.save()


def draw_flow_box(c: canvas.Canvas, text: str, x: int, y: int, width: int, height: int, kind: str = "process") -> None:
    if kind == "start":
        c.roundRect(x, y, width, height, 18, stroke=True, fill=False)
    elif kind == "decision":
        points = [
            x + width / 2,
            y + height,
            x + width,
            y + height / 2,
            x + width / 2,
            y,
            x,
            y + height / 2,
        ]
        c.line(points[0], points[1], points[2], points[3])
        c.line(points[2], points[3], points[4], points[5])
        c.line(points[4], points[5], points[6], points[7])
        c.line(points[6], points[7], points[0], points[1])
    else:
        c.rect(x, y, width, height)

    c.setFont(PDF_FONT_NAME, 9)
    lines = split_pdf_text(text, 28)
    text_y = y + height / 2 + (len(lines) - 1) * 5
    for line in lines:
        c.drawCentredString(x + width / 2, text_y, line)
        text_y -= 12


def split_pdf_text(text: str, max_len: int) -> list[str]:
    words = text.split()
    lines = []
    current = ""

    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= max_len:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word

    if current:
        lines.append(current)

    return lines


def connect_flow(c: canvas.Canvas, start: tuple[int, int], end: tuple[int, int]) -> None:
    c.line(start[0], start[1], end[0], end[1])
    c.circle(end[0], end[1], 2, fill=True)


def generate_flowchart_pdf(output_path: Path) -> None:
    register_pdf_font()
    c = canvas.Canvas(str(output_path), pagesize=landscape(A4))
    c.setTitle("Блок-схема алгоритма приложения")
    c.setFont(PDF_FONT_NAME, 18)
    c.drawString(40, 550, "Алгоритм работы приложения «МебельОрг»")
    c.setFont(PDF_FONT_NAME, 10)
    c.drawString(40, 532, "Оформлено в виде блок-схемы по ГОСТ 19.701-90.")

    c.setStrokeColor(colors.black)
    c.setFillColor(colors.black)
    draw_flow_box(c, "Старт приложения", 65, 455, 150, 45, "start")
    draw_flow_box(c, "Создать БД и импортировать данные при первом запуске", 65, 365, 150, 65)
    draw_flow_box(c, "Показать окно входа", 65, 285, 150, 45)
    draw_flow_box(c, "Пользователь авторизован?", 295, 270, 150, 75, "decision")
    draw_flow_box(c, "Гость: список товаров без поиска", 535, 385, 180, 55)
    draw_flow_box(c, "Клиент: список товаров без поиска", 535, 315, 180, 55)
    draw_flow_box(c, "Менеджер: товары с поиском, фильтром, сортировкой и заказы", 535, 225, 180, 70)
    draw_flow_box(c, "Администратор: CRUD товаров и заказов", 535, 130, 180, 70)
    draw_flow_box(c, "Проверить права роли", 295, 165, 150, 50)
    draw_flow_box(c, "Выход на окно входа", 295, 80, 150, 45)
    draw_flow_box(c, "Конец", 65, 80, 150, 45, "start")

    connect_flow(c, (140, 455), (140, 430))
    connect_flow(c, (140, 365), (140, 330))
    connect_flow(c, (215, 307), (295, 307))
    connect_flow(c, (445, 307), (535, 412))
    connect_flow(c, (445, 307), (535, 342))
    connect_flow(c, (445, 307), (535, 260))
    connect_flow(c, (445, 307), (535, 165))
    connect_flow(c, (370, 270), (370, 215))
    connect_flow(c, (370, 165), (370, 125))
    connect_flow(c, (295, 102), (215, 102))

    c.setFont(PDF_FONT_NAME, 8)
    c.drawString(456, 405, "гость")
    c.drawString(456, 338, "клиент")
    c.drawString(456, 256, "менеджер")
    c.drawString(456, 161, "администратор")
    c.save()


def create_screenshots_docx(output_path: Path, screenshots: list[Path]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Скриншоты корректной работы системы")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("ООО «МебельОрг». Демонстрационный экзамен, вариант 5.")

    captions = [
        "Окно входа при запуске приложения",
        "Гостевой просмотр товаров без фильтрации, сортировки и поиска",
        "Интерфейс менеджера: поиск, фильтр по скидке и сортировка",
        "Интерфейс администратора со средствами управления товарами",
        "Форма добавления и редактирования товара",
        "Экран заказов администратора",
        "Форма добавления и редактирования заказа с позициями",
    ]

    for index, screenshot in enumerate(screenshots):
        document.add_page_break()
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Рисунок {index + 1}. {captions[index]}")
        run.font.bold = True
        run.font.color.rgb = RGBColor(75, 0, 130)
        document.add_picture(str(screenshot), width=Inches(7.0))

    document.save(output_path)


def main() -> None:
    initialize_database()
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)

    create_sql_dump(PROJECT_ROOT / "database" / "mebelorg_dump.sql")
    generate_er_pdf(DOCS_DIR / "ER-diagram.pdf")
    generate_flowchart_pdf(DOCS_DIR / "flowchart.pdf")
    screenshots = generate_screenshots()
    create_screenshots_docx(DOCS_DIR / "screenshots.docx", screenshots)

    print("Artifacts generated")


if __name__ == "__main__":
    main()
''',
    # ============================================================
    # END FILE: tools/generate_artifacts.py
    # ============================================================

    # ============================================================
    # BEGIN FILE: tools/smoke_test.py
    # ============================================================
    'tools/smoke_test.py': r'''import sys
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[1]
APP_DIR = PROJECT_ROOT / "app"
sys.path.insert(0, str(APP_DIR))

from config import DB_ENGINE
from database import get_connection, rebuild_database
from repositories import AuthRepository, LookupRepository, OrderRepository, ProductRepository


def main() -> None:
    rebuild_database()

    expected_counts = {
        "roles": 3,
        "users": 10,
        "products": 10,
        "pickup_points": 36,
        "orders": 10,
        "order_items": 20,
    }

    with get_connection() as connection:
        for table_name, expected in expected_counts.items():
            count = connection.execute(f"SELECT COUNT(*) FROM {table_name}").fetchone()[0]
            print(f"{table_name}: {count}")
            assert count == expected

    auth_repository = AuthRepository()
    admin = auth_repository.authenticate("94d5ous@gmail.com", "uzWC67")
    manager = auth_repository.authenticate("ptec8ym@yahoo.com", "LdNyos")
    client = auth_repository.authenticate("yzls62@outlook.com", "JlFRCZ")

    assert admin and admin.role_name == "Администратор"
    assert manager and manager.role_name == "Менеджер"
    assert client and client.role_name == "Авторизированный клиент"

    product_repository = ProductRepository()
    all_products = product_repository.list_products()
    filtered_products = product_repository.list_products("диван", "15% и более", "Цена", "DESC")
    sorted_by_stock = product_repository.list_products("", "Все диапазоны", "Количество на складе", "ASC")
    zero_stock = [product for product in all_products if product["stock_quantity"] == 0]

    assert len(all_products) == 10
    assert len(filtered_products) == 1
    assert filtered_products[0]["article"] == "F325D4"
    assert sorted_by_stock[0]["stock_quantity"] == 0
    assert zero_stock

    try:
        product_repository.delete_product(all_products[0]["id"])
    except ValueError:
        pass
    else:
        raise AssertionError("Товар из заказа не должен удаляться")

    product_id = product_repository.save_product(
        {
            "article": "TEST01",
            "name": "Тестовый мебельный товар",
            "category_name": all_products[0]["category_name"],
            "manufacturer_name": all_products[0]["manufacturer_name"],
            "supplier_name": all_products[0]["supplier_name"],
            "price": 100.0,
            "unit_name": all_products[0]["unit_name"],
            "stock_quantity": 5,
            "discount": 0.0,
            "description": "Товар для автоматической проверки.",
            "photo_path": None,
        }
    )
    product = product_repository.get_product(product_id)
    assert product and product["article"] == "TEST01"
    product["price"] = 120.0
    product_repository.save_product(product)
    assert product_repository.get_product(product_id)["price"] == 120.0
    product_repository.delete_product(product_id)

    order_repository = OrderRepository()
    orders = order_repository.list_orders()
    assert len(orders) == 10
    assert all(order["items_summary"] for order in orders)

    lookup_repository = LookupRepository()
    pickup_point = lookup_repository.get_pickup_points()[0]
    status = lookup_repository.get_statuses()[0]
    assert order_repository.parse_items_text("А112Т4") == [{"article": "А112Т4", "quantity": 1}]

    order_id = order_repository.save_order(
        {
            "pickup_point_id": pickup_point["id"],
            "client_name": "Тестовый клиент",
            "receive_code": "999",
            "status_id": status["id"],
            "order_date": "2026-06-05",
            "delivery_date": "2026-06-06",
            "items": order_repository.parse_items_text("А112Т4, 1"),
        }
    )
    order = order_repository.get_order(order_id)
    assert order and order["items"][0]["article"] == "А112Т4"
    order_repository.save_order(
        {
            "id": order_id,
            "pickup_point_id": pickup_point["id"],
            "client_name": "Тестовый клиент",
            "receive_code": "999",
            "status_id": status["id"],
            "order_date": "2026-06-05",
            "delivery_date": "2026-06-07",
            "items": order_repository.parse_items_text("G843H5, 2"),
        }
    )
    assert order_repository.get_order(order_id)["delivery_date"] == "2026-06-07"
    order_repository.delete_order(order_id)
    assert order_repository.get_order(order_id) is None

    print(f"SMOKE TEST OK ({DB_ENGINE})")


if __name__ == "__main__":
    main()
''',
    # ============================================================
    # END FILE: tools/smoke_test.py
    # ============================================================

}


def extract_project(target: Path) -> None:
    target.mkdir(parents=True, exist_ok=True)
    for directory in EMPTY_DIRS:
        (target / directory).mkdir(parents=True, exist_ok=True)
    for relative_path, content in FILES.items():
        output_path = target / relative_path
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding='utf-8', newline='')
    print(f'Проект распакован: {target}')


def list_files() -> None:
    for relative_path in sorted(FILES):
        print(relative_path)
    for directory in sorted(EMPTY_DIRS):
        print(directory + '/')


def verify_project(target: Path) -> None:
    missing = [path for path in FILES if not (target / path).exists()]
    if missing:
        raise SystemExit('Не найдены файлы после распаковки: ' + ', '.join(missing))

    python_files = sorted(target.rglob('*.py'))
    for python_file in python_files:
        py_compile.compile(str(python_file), doraise=True)
    print(f'Синтаксис Python проверен: {len(python_files)} файлов')

    smoke_test = target / 'tools' / 'smoke_test.py'
    result = subprocess.run([sys.executable, str(smoke_test)], cwd=str(target), text=True)
    if result.returncode != 0:
        raise SystemExit(result.returncode)
    print('Проверка проекта завершена успешно')


def main() -> None:
    parser = argparse.ArgumentParser(description='Распаковщик исходников МебельОрг')
    parser.add_argument('--list', action='store_true', help='показать файлы внутри пакета')
    parser.add_argument('--extract', metavar='DIR', help='распаковать проект в указанную папку')
    parser.add_argument('--verify', metavar='DIR', help='проверить уже распакованный проект')
    args = parser.parse_args()

    if args.list:
        list_files()
    if args.extract:
        extract_project(Path(args.extract).resolve())
    if args.verify:
        verify_project(Path(args.verify).resolve())
    if not args.list and not args.extract and not args.verify:
        parser.print_help()


if __name__ == '__main__':
    main()
